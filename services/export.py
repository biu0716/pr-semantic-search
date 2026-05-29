"""
导出 service。

把结果文本导出成 docx / md / txt 的二进制内容。
逻辑与原 streamlit_app.py 一致；format 参数改成简单的字符串标识，
方便前端传参（"docx" / "md" / "txt"）。
"""

from io import BytesIO

from docx import Document


def build_docx_bytes(text: str, title: str = "PR Agent 导出结果") -> bytes:
    doc = Document()
    doc.add_heading(title, level=1)

    # 按空行分段，导出到 Word 里会更自然
    blocks = [block.strip() for block in text.split("\n\n") if block.strip()]
    for block in blocks:
        doc.add_paragraph(block)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


# 返回：(二进制内容, MIME 类型, 建议文件名)
def get_export_payload(text: str, export_format: str, title: str = "PR Agent 导出结果"):
    fmt = (export_format or "").lower()

    if fmt in ("docx", "word"):
        return (
            build_docx_bytes(text, title),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "pr_agent_output.docx",
        )

    if fmt in ("md", "markdown"):
        return text.encode("utf-8"), "text/markdown", "pr_agent_output.md"

    return text.encode("utf-8"), "text/plain", "pr_agent_output.txt"
