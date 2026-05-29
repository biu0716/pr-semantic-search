import re
import io
import inspect
import subprocess
from pathlib import Path
import tempfile
from io import BytesIO

import pandas as pd
import streamlit as st
import streamlit.components.v1 as components
from dotenv import load_dotenv
from docx import Document
from markitdown import MarkItDown

load_dotenv()

from src.pr_agent import (
    generate_pr,
    generate_core_pr,
    generate_titles,
    check_text,
    normalize_pr,
    rewrite_text,
    call_llm,
    strip_titles_section,
    clean_title_result,
)
# 设置Streamlit 页面基础配置
# - Page_title: 浏览器标签页标题
# - Page_icon: 浏览器标签页图标
# - Layout="wide"使用宽屏布局，适合左右双栏页面
# - initial_sidebar_state="collapsed": 让侧边栏默认收起
# 这样用户第一眼会先看到主页面里的 “任务选择、输入和结果”，
# 而不是先被“高级功能”吸引注意力
st.set_page_config(
    page_title="PR Agent",
    page_icon="🧠",
    layout="wide",
    initial_sidebar_state="collapsed",
)

st.markdown("""
<style>

/* 上传区：把英文提示替换成中文（主区和侧边栏都会生效） */
div[data-testid="stFileUploaderDropzone"] [data-testid="stMarkdownContainer"] p {
    font-size: 0 !important;
    line-height: 1.5 !important;
    white-space: normal !important;
    position: relative;
    min-height: 22px;
}

div[data-testid="stFileUploaderDropzone"] [data-testid="stMarkdownContainer"] p::before {
    content: "拖拽文件到这里，或点击上传";
    font-size: 14px;
    color: #374151;
    position: absolute;
    left: 0;
    top: 0;
}

/* 上传区小字：改成中文 */
div[data-testid="stFileUploaderDropzone"] small {
    font-size: 0 !important;
    display: block !important;
    position: relative;
    min-height: 18px;
}

div[data-testid="stFileUploaderDropzone"] small::before {
    content: "支持当前任务允许的文件类型";
    font-size: 13px;
    color: #9ca3af;
    position: absolute;
    left: 0;
    top: 0;
}

/* 上传按钮：Browse files -> 选择文件 */
div[data-testid="stFileUploaderDropzone"] button {
    font-size: 0 !important;
    position: relative;
}

div[data-testid="stFileUploaderDropzone"] button::after {
    content: "选择文件";
    font-size: 14px;
    color: #111827;
    position: absolute;
    left: 50%;
    top: 50%;
    transform: translate(-50%, -50%);
    white-space: nowrap;
}

/* 左侧边栏整体
    - background 控制边栏底色
    - border-right 右侧分隔线
    - min-width / max-width 固定边栏宽度，过宽会影响阅读体验 */
[data-testid="stSidebar"] {
    background: #fbfbfb;
    border-right: 1px solid rgba(15, 23, 42, 0.06);
}

/* 只在侧边栏展开时固定宽度 */
[data-testid="stSidebar"][aria-expanded="true"] {
    min-width: 320px !important;
    max-width: 320px !important;
}

/* 主内容区域固定画布宽度，并始终居中 */
[data-testid="stAppViewContainer"] .main .block-container {
    max-width: 1240px;
    margin: 0 auto;
    padding-left: 2rem;
    padding-right: 2rem;
}


/* 侧边栏内部顶部留白 */
[data-testid="stSidebar"] .block-container {
    padding-top: 1.4rem;
}

/* 顶部 Hero 区 （PR Agent 大标题那整块）
    - padding 内部留白
    - border 淡边框
    - border-radius 圆角
    - background 渐变底色
    - box-shadow 阴影 */
.pr-hero {
    padding: 18px 22px;
    border: 1px solid rgba(15, 23, 42, 0.05);
    border-radius: 20px;
    background: #ffffff;
    box-shadow: 0 1px 2px rgba(15, 23, 42, 0.04);
    margin-bottom: 28px;
}
            
/* 顶部很小那行说明文字
    比如 Automotive PR Workflow
    作用是做“产品类别提示”，不要太抢眼 */
.pr-eyebrow {
    font-size: 12px;
    color: #6b7280;
    margin-bottom: 8px;
    letter-spacing: 0.02em;
}


/* 顶部主标题 PR Agent
    - font-size 字体大小
    - font-weight 字体粗细，800 是很粗的标题字重
    - Letter-spacing 字母间距，适当增加可以让标题更有气势
    - line-height 行高，1.1 是比较紧凑的行距 - color 字体颜色，#111827 是接近纯黑的深灰色，视觉上更柔和一些 */
.pr-title {
    font-size: 24px;
    font-weight: 700;
    line-height: 1.12;
    letter-spacing: -0.02em;
    color: #111827;
    margin-bottom: 4px;
}

/* 页面副标题
    解释这个页面是干什么的 */
.pr-subtitle {
    font-size: 14px;
    color: #6b7280;
    line-height: 1.65;
    margin-bottom: 0;
    max-width: 760px;
}

/* 统一步骤标题 */
.pr-step-title {
    font-size: 26px;
    font-weight: 700;
    line-height: 1.18;
    letter-spacing: -0.02em;
    color: #111827;
    margin: 0 0 8px 0;
}


.pr-step-desc {
    font-size: 14px;
    color: #6b7280;
    line-height: 1.7;
    margin-bottom: 22px;
}

/* 更窄屏幕下，把步骤标题自动缩一点，避免右侧换行太难看 */
@media (max-width: 1200px) {
    .pr-step-title {
        font-size: 22px;
    }

    .pr-step-desc {
        margin-bottom: 18px;
    }

    .pr-empty-state {
        min-height: 220px;
        padding: 18px 18px;
    }
}
            
@media (max-width: 900px) {
    .pr-task-choice-card {
        min-height: auto;
    }
            
    .pr-task-choice-desc {
        min-height: 0;
    }
}

    @media (max-width: 1200px) {
        .pr-step-title {
            font-size: 22px;
        }

        .pr-step-desc {
            margin-bottom: 18px;
        }

        .pr-empty-state {
            min-height: 220px;
            padding: 18px 18px;
        }

        .pr-home-inner {
            padding-left: 18px;
            padding-right: 18px;
        }
    }
            

/* 顶部四个统计卡片
    比如“能力模块 / 已支持上传 / 知识库状态 / 适用场景” */            
.pr-stat {
    border: 1px solid rgba(15, 23, 42, 0.05);
    border-radius: 24px;
    padding: 20px 22px;
    background: rgba(255, 255, 255, 0.96);
    box-shadow: 0 6px 18px rgba(15, 23, 42, 0.04);
}

/* 统计卡的小标签文字
    比如”能力模块“这类上方说明 */
.pr-stat-label {
    font-size: 13px;
    color: #6b7280;
    margin-bottom: 8px;
}

/* 统计卡的主要内容
    比如”4 / docx / txt / xlsx / 汽车PR“ */
.pr-stat-value {
    font-size: 17px;
    font-weight: 700;
    color: #111827;
    line-height: 1.35;
}

/* 左右两块主卡片
    左边输入卡、右边结果卡都会用到这个 class */
.pr-card {
    border: 1px solid rgba(15, 23, 42, 0.05);
    border-radius: 28px;
    padding: 24px;
    background: rgba(255, 255, 255, 0.96);
    box-shadow: 0 10px 28px rgba(15, 23, 42, 0.05);
    
    /* 让左边“传播内容生成”和右边“输出结果”顶部卡片高度更接近 */
    min-height: 150px;
}

/* 卡片标题
    比如“传播内容生成”“输出结果” */
.pr-section-title {
    font-size: 24px;
    font-weight: 700;
    letter-spacing: -0.01em;
    color: #111827;
    margin-bottom: 6px;
}

/* 卡片标题下的说明文字 */
.pr-section-desc {
    font-size: 14px;
    color: #6b7280;
    line-height: 1.5;
    margin-bottom: 16px;
}

/* 结果区默认空态提示卡
    用浅灰底替代 st.info 的蓝底，并给它固定高度
    让它和左边输入框视觉上更对齐 */
.pr-empty-state {
    border: 1px solid rgba(15, 23, 42, 0.05);
    border-radius: 18px;
    background: #f6f7fb;
    color: #475467;
    padding: 24px 24px;
    font-size: 15px;
    line-height: 1.7;
    min-height: 300px;
    display: flex;
    align-items: flex-start;
}

.pr-empty-state-title {
    font-size: 18px;
    font-weight: 700;
    color: #111827;
    margin-bottom: 10px;
}

.pr-empty-state-desc {
    font-size: 15px;
    line-height: 1.8;
    color: #475467;
} 

.pr-result-state-card {
    border: 1px solid rgba(15, 23, 42, 0.06);
    border-radius: 18px;
    background: #fafafc;
    padding: 18px 18px;
    min-height: 148px;
    margin-bottom: 14px;
}

.pr-result-mode-badge {
    display: inline-block;
    font-size: 12px;
    font-weight: 700;
    color: #475467;
    background: #f3f4f6;
    border-radius: 999px;
    padding: 5px 10px;
    margin-bottom: 12px;
}

.pr-result-state-title {
    font-size: 18px;
    font-weight: 700;
    color: #111827;
    margin-bottom: 8px;
    line-height: 1.4;
}

.pr-result-state-desc {
    font-size: 14px;
    line-height: 1.75;
    color: #6b7280;
}

.pr-result-scroll {
    max-height: 560px;
    overflow-y: auto;
    padding-right: 4px;
}
                 
.pr-task-guide {
    border: 1px solid rgba(15, 23, 42, 0.05);
    border-radius: 20px;
    background: #f8fafc;
    padding: 18px 20px;
    margin-bottom: 16px;
}

.pr-task-guide-title {
    font-size: 15px;
    font-weight: 700;
    color: #111827;
    margin-bottom: 6px;
}

.pr-task-guide-desc {
    font-size: 13px;
    line-height: 1.6;
    color: #6b7280;
    margin-bottom: 10px;
}

.pr-task-guide ul {
    margin: 0;
    padding-left: 18px;
}

.pr-task-guide li {
    font-size: 14px;
    line-height: 1.7;
    color: #374151;
    margin-bottom: 6px;
}
/* Streamlit 按钮样式
    - border-radius 圆角
    - height 按钮高度
    - font-weight 字体加粗 */
/* 普通按钮：改成浅色，更轻一点 */
div.stButton > button[kind="secondary"] {
    border-radius: 18px;
    height: 54px;
    font-weight: 700;
    border: 1px solid rgba(15, 23, 42, 0.08);
    background: #ffffff;
    color: #111827;
    box-shadow: 0 4px 12px rgba(15, 23, 42, 0.05);
}

/* 普通按钮悬停时 */
div.stButton > button[kind="secondary"]:hover {
    background: #f8fafc;
    border-color: rgba(15, 23, 42, 0.12);
    color: #111827;
    box-shadow: 0 6px 16px rgba(15, 23, 42, 0.07);
}

/* 主按钮：改成白色 */
            
div.stButton > button[kind="primary"] {
    border-radius: 16px;
    height: 48px;
    font-weight: 600;
    border: 1px solid rgba(15, 23, 42, 0.10);
    background: #1f2937;
    color: white;
    box-shadow: 0 4px 12px rgba(15, 23, 42, 0.10);
}

div.stButton > button[kind="primary"]:hover {
    background: #111827;
    border-color: rgba(17, 24, 39, 0.16);
    color: white;
}

div.stButton > button[kind="primary"]:focus,
div.stButton > button[kind="primary"]:focus-visible,
div.stButton > button[kind="primary"]:active {
    background: #1d293b;
    border-color: rgba(31, 41, 55, 0.16);
    color: white;
    box-shadow: 0 0 0 2px rgba(15, 23, 42, 0.04);
    outline: none;
}

div.stButton > button {
    white-space: nowrap !important;
}

div.stButton > button[kind="secondary"] {
    min-width: 88px;
}

div.stButton > button[kind="primary"] {
    min-width: 120px;
}
            
/* 文本输入框：默认状态 */
div[data-testid="stTextArea"] textarea {
    border-radius: 18px !important;
    border: 1px solid rgba(15, 23, 42, 0.06) !important;
    background: #fcfcfd !important;
    box-shadow: none !important;
    outline: none !important;
    color: #111827 !important;
}

/* 外层容器也一起控制 */
div[data-testid="stTextArea"] > div {
    border-radius: 18px !important;
    border: 1px solid rgba(15, 23, 42, 0.06) !important;
    background: #fcfcfd !important;
    box-shadow: none !important;
    outline: none !important;
}

/* 真正选中时，不要红框，改成浅灰蓝 */
div[data-testid="stTextArea"] > div:focus-within {
    border: 1px solid rgba(15, 23, 42, 0.12) !important;
    box-shadow: 0 0 0 2px rgba(15, 23, 42, 0.04) !important;
    outline: none !important;
}

/* 文本框本身获得焦点时，也压掉默认红色 */
div[data-testid="stTextArea"] textarea:focus,
div[data-testid="stTextArea"] textarea:focus-visible {
    border-color: rgba(15, 23, 42, 0.12) !important;
    box-shadow: none !important;
    outline: none !important;
}
            
/* 运行状态提示 */
 .pr-run-status {
     font-size: 14px;
     color: #475467;
     margin-top: 6px;
     margin-bottom: 8px;
}
/* 让进度条上方和按钮之间更舒服一点 */
 .pr-run-gap {
     height: 10px
}
/* 选择模式：外层输入框 */
div[data-baseweb="select"] > div {
    border-radius: 14px !important;
    border: 1px solid rgba(15, 23, 42, 0.08) !important;
    background: #ffffff !important;
    box-shadow: none !important;
    transition: all 0.2s ease;
}

/* 选择模式：点击或聚焦时，不要红色，改成浅灰蓝 */
div[data-baseweb="select"] > div:focus-within {
    border-color: rgba(15, 23, 42, 0.14) !important;
    box-shadow: 0 0 0 2px rgba(15, 23, 42, 0.04) !important;
    outline: none !important;
}

/* 下拉菜单里当前选中的项，不要红底，改成浅灰底 */
div[role="option"][aria-selected="true"] {
    background: #f3f4f6 !important;
    color: #111827 !important;
}
            
/* 下拉菜单鼠标移上去时的颜色 */
div [role="option"]:hover {
    background: #f8fafc !important;
    color: #111827 !important;
}

/* Top K 输入框里的数字输入区域 */
div[data-testid="stNumberInput"] input{
    border-radius: 14px !important;
    border: 1px solid rgba(15, 23, 42, 0.08) !important;
    background: #ffffff !important;
    box-shadow: none !important;
}

/* Top K 右边的加减号按钮：默认状态 */
div[data-testid="stNumberInput"] button {
    background: transparent !important;
    color: #111827 !important;
    border: none !important;
    box-shadow: none !important;
    outline: none !important;
}
/* Top K 加减号按钮：鼠标移上去时 */
div[data-testid="stNumberInput"] button:hover {
    background: #f3f4f6 !important;
    color: #111827 !important;
}

/* Top K 加减号按钮：点击/聚焦时，不要红色 */
div[data-testid="stNumberInput"] button:active,
div[data-testid="stNumberInput"] button:focus,
div[data-testid="stNumberInput"] button:focus-visible {
    background: #eef2f7 !important;
    color: #111827 !important;
    box-shadow: none !important;
    outline: none !important;
}           
/* 文件上传框 */
div[data-testid="stFileUploader"] section {
    border-radius: 18px;
    border: 1px solid rgba(15, 23, 42, 0.06);
    background: #fcfcfd;            
} 
/* Top K 整个输入组件的外层容器 */
div[data-testid="stNumberInput"] > div {
    border-radius: 16px !important;
    border: 1px solid rgba(15, 23, 42, 0.08) !important;
    background: #ffffff !important;
    box-shadow: none !important;
    overflow: hidden !important;
}

/* Top K 获得焦点时，去掉默认红色边框和红色光晕 */
div[data-testid="stNumberInput"] > div:focus-within {
    border: 1px solid rgba(15, 23, 42, 0.12) !important;
    box-shadow: 0 0 0 2px rgba(15, 23, 42, 0.04) !important;
    outline: none !important;
}

/* Top K 里面真正的输入框 */
div[data-testid="stNumberInput"] input {
    border: none !important;
    box-shadow: none !important;
    background: transparent !important;
    color: #111827 !important;
}

/* Top K 输入框本身获得焦点时，也不要红色 */
div[data-testid="stNumberInput"] input:focus,
div[data-testid="stNumberInput"] input:focus-visible {
    outline: none !important;
    box-shadow: none !important;
    border: none !important;
}

/* Top K 右边的 + / - 按钮 */
div[data-testid="stNumberInput"] button {
    background: transparent !important;
    border: none !important;
    color: #111827 !important;
    box-shadow: none !important;
    outline: none !important;
}

/* + / - 鼠标移上去时 */
div[data-testid="stNumberInput"] button:hover {
    background: #f3f4f6 !important;
    color: #111827 !important;
}

/* + / - 被点击或聚焦时，不要红色 */
div[data-testid="stNumberInput"] button:active,
div[data-testid="stNumberInput"] button:focus,
div[data-testid="stNumberInput"] button:focus-visible {
    background: #eef2f7 !important;
    border: none !important;
    box-shadow: none !important;
    outline: none !important;
    color: #111827 !important;
}

/* 选择模式下拉框外层 */
div[data-baseweb="select"] > div {
    border-radius: 14px !important;
    border: 1px solid rgba(15, 23, 42, 0.08) !important;
    background: #ffffff !important;
    box-shadow: none !important;
}

/* 选择模式下拉框聚焦时，不要红色 */
div[data-baseweb="select"] > div:focus-within {
    border: 1px solid rgba(15, 23, 42, 0.12) !important;
    box-shadow: 0 0 0 2px rgba(15, 23, 42, 0.04) !important;
    outline: none !important;
}
/* 页面中的分割线  如果需要在内容中分割不同部分，可以使用 <hr> 标签，这里设置了它的上下间距 */
hr {
    margin-top: 1.2rem !important !important;
    margin-bottom: 1.2rem !important;
    border-color: rgba(15, 23, 42, 0.06);
}

/* 输出结果正文样式 */
.pr-result-html {
    color: #111827;
    padding-top: 4px;
}

.pr-result-html h3 {
    margin: 18px 0 10px 0;
    font-size: 18px;
    line-height: 1.4;
    font-weight: 700;
    color: #111827;
}

.pr-result-html p {
    margin: 0 0 12px 0;
    font-size: 16px;
    line-height: 1.9;
    color: #374151;
}

.pr-result-html ul {
    margin: 4px 0 14px 22px;
    padding-left: 18px;
}

.pr-result-html li {
    margin-bottom: 10px;
    font-size: 16px;
    line-height: 1.9;
    color: #374151;
}
            
.pr-stage-card {
    border: 1px solid rgba(15, 23, 42, 0.06);
    border-radius: 22px;
    background: #ffffff;
    padding: 18px 20px;
    margin-bottom: 16px;
    box-shadow: 0 6px 18px rgba(15, 23, 42, 0.04);
}

.pr-stage-label {
    display: inline-block;
    font-size: 12px;
    font-weight: 700;
    color: #475467;
    background: #f3f4f6;
    border-radius: 999px;
    padding: 6px 10px;
    margin-bottom: 12px;
}

.pr-stage-title {
    font-size: 20px;
    font-weight: 700;
    color: #111827;
    margin-bottom: 10px;
}

.pr-stage-desc {
    font-size: 14px;
    line-height: 1.8;
    color: #475467;
}

.pr-subsection-title {
    font-size: 16px;
    font-weight: 700;
    color: #111827;
    margin: 18px 0 10px 0;
}

.pr-divider {
    height: 1px;
    background: rgba(15, 23, 42, 0.06);
    margin: 16px 0;
}

/* ==============================
   任务选择卡片样式
   目的：
   1. 让“选择任务”更像工作台入口，而不是表单
   2. 让当前选中的任务更明显
   ============================== */

/* 任务选择卡片
   这里额外加了 margin-bottom，
   目的就是让卡片和下面的”已选择 / 选择某任务”按钮拉开距离，
   不要看起来贴得太近。 */
.pr-task-choice-card {
    border: 1px solid rgba(15, 23, 42, 0.08);
    border-radius: 22px;
    background: #ffffff;
    padding: 22px 20px;
    min-height: 172px;
    margin-bottom: 16px;
    box-shadow: 0 4px 14px rgba(15, 23, 42, 0.04);
    transition: all 0.2s ease;
    display: flex;
    flex-direction: column;
    justify-content: flex-start;
}


.pr-task-choice-card.is-selected {
    border: 1px solid rgba(15, 23, 42, 0.16);
    background: #f8fafc;
    box-shadow: 0 8px 20px rgba(15, 23, 42, 0.06);
}

.pr-task-choice-badge {
    display: inline-block;
    font-size: 12px;
    font-weight: 700;
    color: #475467;
    background: #f3f4f6;
    border-radius: 999px;
    padding: 5px 10px;
    margin-bottom: 12px;
}

.pr-task-choice-card.is-selected .pr-task-choice-badge {
    color: #111827;
    background: #e5e7eb;
}

.pr-task-choice-title {
    font-size: 19px;
    font-weight: 700;
    color: #111827;
    line-height: 1.4;
    margin-bottom: 8px;
    min-height: 28px;
}

.pr-task-choice-desc {
    font-size: 14px;
    line-height: 1.7;
    color: #6b7280;
    min-height: 72px;
}
            
.pr-home-actions {
    margin-top: 12px;
}

/* 让任务卡本身变成一个可点击链接
   目的：
   1. 点击整张卡，而不是点下面单独的按钮
   2. 保留卡片外观，不显示默认蓝色下划线 */

/* ==============================
   当前任务说明卡
   目的：
   1. 不把大段说明塞进按钮里
   2. 让任务按钮保持简洁，说明文字单独展示
   ============================== */

.pr-task-summary-card {
    border: 1px solid rgba(15, 23, 42, 0.06);
    border-radius: 20px;
    background: #fafafc;
    padding: 18px 20px;
    margin-top: 14px;
}

.pr-task-summary-label {
    display: inline-block;
    font-size: 12px;
    font-weight: 700;
    color: #475467;
    background: #f3f4f6;
    border-radius: 999px;
    padding: 5px 10px;
    margin-bottom: 10px;
}

.pr-task-summary-title {
    font-size: 18px;
    font-weight: 700;
    color: #111827;
    margin-bottom: 6px;
    line-height: 1.4;
}

.pr-task-summary-desc {
    font-size: 14px;
    line-height: 1.75;
    color: #6b7280;
}
            
div.stButton > button[kind="secondary"] {
    border-radius: 14px;
    height: 42px;
    font-weight: 600;
    border: 1px solid rgba(15, 23, 42, 0.06);
    background: #f8fafc;
    color: #475467;
    box-shadow: none;
}

div.stButton > button[kind="secondary"]:hover {
    background: #f1f5f9;
    border-color: rgba(15, 23, 42, 0.08);
    color: #111827;
}


/* ==============================
   使用助手：桌面 App 风格悬浮入口
   ============================== */

.pr-assistant-highlight {
    border: 1px solid rgba(15, 23, 42, 0.08);
    border-radius: 18px;
    background: linear-gradient(180deg, #ffffff 0%, #f8fafc 100%);
    box-shadow: 0 6px 18px rgba(15, 23, 42, 0.05);
    padding: 14px 16px;
    margin-bottom: 12px;
}

.pr-assistant-shell {
    border: 1px solid rgba(15, 23, 42, 0.06);
    border-radius: 22px;
    background: #ffffff;
    box-shadow: 0 10px 28px rgba(15, 23, 42, 0.05);
    padding: 20px 20px 18px 20px;
    margin-top: 2px;
    margin-bottom: 12px;
}

.pr-assistant-title {
    font-size: 20px;
    font-weight: 700;
    color: #111827;
    margin-bottom: 6px;
    line-height: 1.4;
}

.pr-assistant-desc {
    font-size: 14px;
    color: #6b7280;
    line-height: 1.75;
    margin-bottom: 0;
}

div[data-testid="stPopover"] {
    position: fixed;
    right: 24px;
    bottom: 24px;
    z-index: 99999;
    user-select: none;
}

div[data-testid="stPopover"] > div > button,
div[data-testid="stPopover"] button {
    width: 62px !important;
    height: 62px !important;
    min-width: 62px !important;
    border-radius: 999px !important;
    background: linear-gradient(180deg, #111827 0%, #0f172a 100%) !important;
    color: #ffffff !important;
    border: none !important;
    box-shadow:
        0 16px 36px rgba(15, 23, 42, 0.20),
        0 4px 10px rgba(15, 23, 42, 0.10) !important;
    font-size: 28px !important;
    font-weight: 700 !important;
    padding: 0 !important;
    transition:
        transform 0.18s ease,
        box-shadow 0.18s ease,
        background 0.18s ease,
        opacity 0.18s ease !important;
    cursor: grab !important;
}

div[data-testid="stPopover"] > div > button:hover,
div[data-testid="stPopover"] button:hover {
    background: linear-gradient(180deg, #0f172a 0%, #111827 100%) !important;
    transform: translateY(-2px) scale(1.02);
    box-shadow:
        0 20px 42px rgba(15, 23, 42, 0.24),
        0 6px 14px rgba(15, 23, 42, 0.12) !important;
}

div[data-testid="stPopover"] > div > button:active,
div[data-testid="stPopover"] button:active {
    cursor: grabbing !important;
    transform: scale(0.98);
}

div[data-testid="stPopover"] > div > button p,
div[data-testid="stPopover"] button p {
    color: #ffffff !important;
    font-size: 28px !important;
    margin: 0 !important;
    line-height: 1 !important;
}

div[data-testid="stPopoverContent"] {
    width: min(460px, calc(100vw - 32px)) !important;
    max-height: 72vh !important;
    overflow-y: auto !important;
    border-radius: 28px !important;
    border: 1px solid rgba(15, 23, 42, 0.06) !important;
    box-shadow:
        0 20px 54px rgba(15, 23, 42, 0.18),
        0 6px 18px rgba(15, 23, 42, 0.08) !important;
    background: rgba(255, 255, 255, 0.98) !important;
    backdrop-filter: blur(14px);
}

@media (max-width: 900px) {
    div[data-testid="stPopover"] {
        right: 16px;
        bottom: 16px;
    }

    div[data-testid="stPopover"] > div > button,
    div[data-testid="stPopover"] button {
        width: 54px !important;
        height: 54px !important;
        min-width: 54px !important;
        font-size: 24px !important;
    }

    div[data-testid="stPopover"] > div > button p,
    div[data-testid="stPopover"] button p {
        font-size: 24px !important;
    }

    div[data-testid="stPopoverContent"] {
        width: min(440px, calc(100vw - 24px)) !important;
    }
}

</style>
""", unsafe_allow_html=True)

# ==============================
# Paths
# ==============================

PROJECT_ROOT = Path(__file__).resolve().parent
from pathlib import Path
import tempfile

UPLOAD_DIR = Path(tempfile.gettempdir()) / "pr_agent_uploads"
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)


# ==============================
# Task-first UX Config
# ==============================

TASK_GROUP_LABELS = {
    "create": "生成新内容",
    "refine": "处理现有稿件",
}

TASK_GROUP_DESCRIPTIONS = {
    "create": "适合从需求出发，生成传播提案、标题、导语或社媒文案。",
    "refine": "适合对已有稿件进行审稿、规范化或改写。",
}

MODE_LABELS = {
    "generate": "生成传播提案 / 标题 / 导语 / 社媒文案",
    "check": "审稿并给修改建议",
    "normalize": "统一术语和官方表达",
    "rewrite": "改写为更正式的 PR 风格",
}

MODE_DESCRIPTIONS = {
    "generate": "基于知识库或上传资料，生成传播主题、传播方向、媒体标题、新闻稿段落和社媒文案。",
    "check": "检查品牌口径、时间表达、风险措辞、术语写法、格式和 PR 语气问题，并给出修改建议。",
    "normalize": "将口语、非官方写法或不规范表述调整为更正式的品牌公关表达。",
    "rewrite": "保留原意，对原文进行更像正式汽车公关稿件的改写优化。",
}

MODE_QUICK_TIPS = {
    "generate": "先描述你想产出什么；有资料时建议上传资料，没有资料也可以直接生成。",
    "check": "可以直接粘贴稿件，也可以上传稿件；系统会更偏向指出问题和修改建议。",
    "normalize": "适合统一术语、品牌写法和更正式的表达口径。",
    "rewrite": "适合在不改变原意的前提下，把文本改得更像正式汽车公关稿件。",
}

GROUP_TO_MODES = {
    "create": ["generate"],
    "refine": ["check", "normalize", "rewrite"],
}

TASK_GROUP_OPTIONS = [
    TASK_GROUP_LABELS["create"],
    TASK_GROUP_LABELS["refine"],
]

TASK_GROUP_LABEL_TO_KEY = {v: k for k, v in TASK_GROUP_LABELS.items()}
LABEL_TO_MODE = {v: k for k, v in MODE_LABELS.items()}

MODE_INPUT_CONFIG = {
    "generate": {
        "field_label": "传播需求",
        "placeholder": "例如：为全新纯电VLE输出传播主题、三个传播方向和五个媒体标题；面向中国媒体，避免卖点罗列。",
        "help": "建议写清：想产出什么、面向谁、强调什么、避免什么。",
        "sample": "为全新纯电VLE输出传播主题、三个传播方向和五个媒体标题，面向中国媒体，避免把传播方向写成卖点分类。",
    },
    "check": {
        "field_label": "待审稿文本",
        "placeholder": "请粘贴待审稿文本，或上传稿件。",
        "help": "适合检查品牌口径、时间表达、风险措辞、术语写法和 PR 语气。",
        "sample": "奔驰将在德国时间19:15发布新车，这款车非常厉害，重新定义豪华mpv体验，是行业最强产品之一。",
    },
    "normalize": {
        "field_label": "待规范化文本",
        "placeholder": "请粘贴待规范化文本，或上传稿件。",
        "help": "适合统一术语、品牌写法和更正式的表达口径。",
        "sample": "奔驰将在德国时间19:15发布新车",
    },
    "rewrite": {
        "field_label": "待改写文本",
        "placeholder": "请粘贴待改写文本，或上传稿件。",
        "help": "适合在不改变原意的前提下，把文本改得更像正式汽车 PR 稿件。",
        "sample": "这款车空间很大，而且挺适合家庭出行",
    },
}

MODE_SHORT_HINTS = {
    "generate": "输入传播需求；有资料可上传，没有资料也可以直接生成。",
    "check": "粘贴原文或上传稿件，系统会输出问题和修改建议。",
    "normalize": "粘贴原文或上传稿件，系统会统一术语和官方表达。",
    "rewrite": "粘贴原文或上传稿件，系统会保留原意并改得更正式。",
}

MODE_OUTPUT_HINTS = {
    "generate": "传播提案、传播方向、标题、导语或社媒文案",
    "check": "问题清单 + 修改建议 + 可参考改法",
    "normalize": "统一后的正式表达",
    "rewrite": "更像正式汽车 PR 稿件的改写结果",
}

REFINE_INSTRUCTION_CONFIG = {
    "check": {
        "field_label": "补充审稿要求（可选）",
        "placeholder": "例如：重点检查品牌口径、时间表达和风险措辞；只输出问题清单，不要整段重写。",
        "help": "上传稿件后，这里用于补充个性化要求，不再需要粘贴全文。",
    },
    "normalize": {
        "field_label": "补充规范化要求（可选）",
        "placeholder": "例如：统一为梅赛德斯-奔驰官方口径；保留原段落结构；不要改写成新闻稿。",
        "help": "上传稿件后，这里用于补充规范化要求。",
    },
    "rewrite": {
        "field_label": "补充改写要求（可选）",
        "placeholder": "例如：改得更像正式汽车 PR 稿件；语气更克制；保留原意，不新增事实。",
        "help": "上传稿件后，这里用于补充改写要求。",
    },
}

def get_refine_input_config(mode: str, has_uploaded_file: bool):
    if has_uploaded_file and mode in REFINE_INSTRUCTION_CONFIG:
        merged = MODE_INPUT_CONFIG[mode].copy()
        merged.update(REFINE_INSTRUCTION_CONFIG[mode])
        return merged
    return MODE_INPUT_CONFIG[mode]


USAGE_ASSISTANT_QUICK_ANSWERS = {
    "上传后还要不要再贴全文？": "如果已经上传稿件，通常不需要再粘贴全文。下面的输入框此时是给你补充个性化处理要求用的，比如“重点检查品牌口径”或“保留原段落结构”。",
    "审稿 / 规范化 / 改写怎么选？": "审稿适合找问题、给建议；规范化适合统一术语和官方口径；改写适合保留原意、把文字改得更像正式汽车 PR 稿件。",
    "为什么结果不够像我要的？": "通常是因为要求写得还不够具体。你可以补充：面向谁、强调什么、避免什么、是否允许重写结构、是否允许新增结论。",
    "这个工具更适合什么场景？": "它更适合汽车 PR 的真实工作流，比如基于已有稿件做审稿、改写、统一术语，或者基于产品资料生成传播提案，而不是泛泛聊天。",
}

def answer_usage_assistant_question(
    mode: str,
    question: str,
    has_uploaded_file: bool,
) -> str:
    question = (question or "").strip()

    if not question:
        return ""

    if question in USAGE_ASSISTANT_QUICK_ANSWERS:
        if question == "上传后还要不要再贴全文？" and not has_uploaded_file:
            return "你现在还没上传稿件，所以输入框仍然是给你粘贴正文用的。只有上传稿件后，这个输入框才会切换为“补充个性化处理要求”。"
        return USAGE_ASSISTANT_QUICK_ANSWERS[question]

    prompt = f"""
你是 PR Agent 的产品使用助手，不负责完成公关稿件内容，只负责回答“这个产品怎么用”。

【产品定位】
- 面向汽车公关（PR）工作流
- 不是通用聊天机器人
- 重点是降低学习成本，帮助用户快速完成任务

【当前页面状态】
- 当前模式：{MODE_LABELS.get(mode, mode)}
- 是否已上传稿件：{"是" if has_uploaded_file else "否"}

【回答要求】
1. 只回答“怎么使用这个产品”。
2. 不要替用户完成传播内容、审稿或改写任务本身。
3. 优先告诉用户下一步该点什么、填什么。
4. 如果当前已上传稿件，要明确提醒：下面输入框用于补充个性化处理要求，不必再粘贴全文。
5. 不要编造当前产品里没有实现的功能。
6. 用中文回答，控制在 120 字以内，清楚、直接、像产品内说明。

【用户问题】
{question}
""".strip()

    return call_llm(prompt).strip()

def render_usage_assistant_content(
    mode: str | None,
    has_uploaded_file: bool,
):
    effective_mode = mode or "generate"
    title_text = "产品使用助手"
    desc_text = "已为你自动展开。点击圆球可收起；需要时再点开即可。"

    if has_uploaded_file:
        desc_text = "你已上传文件。这里适合问：接下来要不要再贴全文、补充要求该怎么写。点击右下角圆球可收起。"

    st.markdown(
        f"""
        <div class="pr-assistant-shell">
            <div class="pr-assistant-title">🤖 {title_text}</div>
            <div class="pr-assistant-desc">{desc_text}</div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    st.markdown(
        """
        <div class="pr-assistant-highlight">
            <div style="font-size:13px; font-weight:700; color:#111827; margin-bottom:6px;">
                快速说明
            </div>
            <div style="font-size:13px; line-height:1.7; color:#475467;">
                点下面的问题可以快速上手；也可以直接输入你对这个产品使用方式的疑问。
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    answer_placeholder = st.empty()

    answer = st.session_state.get("usage_assistant_answer", "").strip()
    if answer:
        answer_placeholder.markdown(
            f"""
            <div style="
                border:1px solid rgba(15, 23, 42, 0.08);
                border-radius:16px;
                background:#ffffff;
                padding:14px 16px;
                margin-top:4px;
                margin-bottom:12px;
                box-shadow: 0 4px 12px rgba(15, 23, 42, 0.04);
            ">
                <div style="font-size:13px; font-weight:700; color:#111827; margin-bottom:8px;">
                    使用助手回答
                </div>
                <div style="font-size:14px; line-height:1.75; color:#374151;">
                    {html.escape(answer)}
                </div>
            </div>
            """,
            unsafe_allow_html=True,
        )

    quick_col1, quick_col2 = st.columns(2, gap="small")

    with quick_col1:
        if st.button(
            "上传后还要不要再贴全文？",
            key=f"help_q1_{effective_mode}",
            use_container_width=True,
        ):
            st.session_state.usage_assistant_answer = answer_usage_assistant_question(
                mode=effective_mode,
                question="上传后还要不要再贴全文？",
                has_uploaded_file=has_uploaded_file,
            )
            st.rerun()

        if st.button(
            "为什么结果不够像我要的？",
            key=f"help_q3_{effective_mode}",
            use_container_width=True,
        ):
            st.session_state.usage_assistant_answer = answer_usage_assistant_question(
                mode=effective_mode,
                question="为什么结果不够像我要的？",
                has_uploaded_file=has_uploaded_file,
            )
            st.rerun()

    with quick_col2:
        if st.button(
            "审稿 / 规范化 / 改写怎么选？",
            key=f"help_q2_{effective_mode}",
            use_container_width=True,
        ):
            st.session_state.usage_assistant_answer = answer_usage_assistant_question(
                mode=effective_mode,
                question="审稿 / 规范化 / 改写怎么选？",
                has_uploaded_file=has_uploaded_file,
            )
            st.rerun()

        if st.button(
            "这个工具更适合什么场景？",
            key=f"help_q4_{effective_mode}",
            use_container_width=True,
        ):
            st.session_state.usage_assistant_answer = answer_usage_assistant_question(
                mode=effective_mode,
                question="这个工具更适合什么场景？",
                has_uploaded_file=has_uploaded_file,
            )
            st.rerun()

    custom_question = st.text_input(
        "也可以直接问",
        placeholder="例如：我上传了稿件以后，这里还需要输入什么？",
        key=f"usage_assistant_input_{effective_mode}",
    )

    if st.button(
        "提问",
        key=f"usage_assistant_submit_{effective_mode}",
        use_container_width=True,
    ):
        st.session_state.usage_assistant_answer = answer_usage_assistant_question(
            mode=effective_mode,
            question=custom_question,
            has_uploaded_file=has_uploaded_file,
        )
        st.rerun()

def inject_desktop_assistant_js():
    components.html(
        """
        <script>
        const rootWindow = window.parent || window;
        const rootDocument = rootWindow.document;

        const STORAGE_KEY = "pr_agent_desktop_assistant_pos_v4";
        const AUTO_OPEN_KEY = "pr_agent_assistant_auto_opened_v3";

        function clamp(value, min, max) {
            return Math.min(Math.max(value, min), max);
        }

        function getLocalStorage() {
            try {
                return rootWindow.localStorage;
            } catch (e) {
                return null;
            }
        }

        function getSessionStorage() {
            try {
                return rootWindow.sessionStorage;
            } catch (e) {
                return null;
            }
        }

        function savePosition(left, top) {
            const storage = getLocalStorage();
            if (!storage) return;

            try {
                storage.setItem(
                    STORAGE_KEY,
                    JSON.stringify({ left, top })
                );
            } catch (e) {}
        }

        function loadPosition() {
            const storage = getLocalStorage();
            if (!storage) return null;

            try {
                const raw = storage.getItem(STORAGE_KEY);
                if (!raw) return null;
                return JSON.parse(raw);
            } catch (e) {
                return null;
            }
        }

        function isPopoverOpen() {
            return !!rootDocument.querySelector('div[data-testid="stPopoverContent"]');
        }

        function autoOpenOnce(buttonEl) {
            const storage = getSessionStorage();
            if (!storage) return;
            if (storage.getItem(AUTO_OPEN_KEY) === "1") return;

            storage.setItem(AUTO_OPEN_KEY, "1");

            rootWindow.setTimeout(() => {
                if (!isPopoverOpen()) {
                    buttonEl.click();
                }
            }, 420);
        }

        function resetToDefaultPosition(popoverEl) {
            const gap = rootWindow.innerWidth <= 900 ? 16 : 24;
            const buttonRect = popoverEl.getBoundingClientRect();

            const left = rootWindow.innerWidth - buttonRect.width - gap;
            const top = rootWindow.innerHeight - buttonRect.height - gap;

            popoverEl.style.right = "auto";
            popoverEl.style.bottom = "auto";
            popoverEl.style.left = left + "px";
            popoverEl.style.top = top + "px";

            savePosition(left, top);
        }

        function attachDesktopBehavior() {
            const popoverEl = rootDocument.querySelector('div[data-testid="stPopover"]');
            if (!popoverEl) return;

            const buttonEl =
                popoverEl.querySelector(':scope > div > button') ||
                popoverEl.querySelector("button");

            if (!buttonEl) return;

            if (popoverEl.dataset.desktopReady === "1") {
                autoOpenOnce(buttonEl);
                return;
            }

            popoverEl.dataset.desktopReady = "1";

            popoverEl.style.position = "fixed";
            popoverEl.style.right = "24px";
            popoverEl.style.bottom = "24px";
            popoverEl.style.left = "auto";
            popoverEl.style.top = "auto";
            popoverEl.style.zIndex = "99999";
            popoverEl.style.transition = "left 0.18s ease, top 0.18s ease";
            popoverEl.style.userSelect = "none";

            buttonEl.title = "点击展开或收起使用助手；拖动可移动位置";

            const saved = loadPosition();
            if (saved && typeof saved.left === "number" && typeof saved.top === "number") {
                popoverEl.style.right = "auto";
                popoverEl.style.bottom = "auto";
                popoverEl.style.left = saved.left + "px";
                popoverEl.style.top = saved.top + "px";
            }

            let dragging = false;
            let moved = false;
            let startX = 0;
            let startY = 0;
            let startLeft = 0;
            let startTop = 0;

            function onMove(event) {
                if (!dragging) return;

                const dx = event.clientX - startX;
                const dy = event.clientY - startY;

                if (Math.abs(dx) > 3 || Math.abs(dy) > 3) {
                    moved = true;
                }

                const buttonRect = buttonEl.getBoundingClientRect();
                const gap = 8;

                const nextLeft = clamp(
                    startLeft + dx,
                    gap,
                    rootWindow.innerWidth - buttonRect.width - gap
                );

                const nextTop = clamp(
                    startTop + dy,
                    gap,
                    rootWindow.innerHeight - buttonRect.height - gap
                );

                popoverEl.style.right = "auto";
                popoverEl.style.bottom = "auto";
                popoverEl.style.left = nextLeft + "px";
                popoverEl.style.top = nextTop + "px";
            }

            function onUp() {
                if (!dragging) return;

                dragging = false;
                buttonEl.style.cursor = "grab";
                popoverEl.style.transition = "left 0.18s ease, top 0.18s ease";

                const rect = popoverEl.getBoundingClientRect();
                savePosition(rect.left, rect.top);

                rootWindow.removeEventListener("pointermove", onMove);
                rootWindow.removeEventListener("pointerup", onUp);

                if (moved) {
                    buttonEl.dataset.skipClickOnce = "1";
                    rootWindow.setTimeout(() => {
                        buttonEl.dataset.skipClickOnce = "0";
                    }, 140);
                }
            }

            buttonEl.addEventListener("pointerdown", (event) => {
                dragging = true;
                moved = false;
                startX = event.clientX;
                startY = event.clientY;

                const rect = popoverEl.getBoundingClientRect();
                startLeft = rect.left;
                startTop = rect.top;

                popoverEl.style.transition = "none";
                buttonEl.style.cursor = "grabbing";

                rootWindow.addEventListener("pointermove", onMove);
                rootWindow.addEventListener("pointerup", onUp);
            });

            buttonEl.addEventListener("dblclick", (event) => {
                event.preventDefault();
                event.stopPropagation();
                resetToDefaultPosition(popoverEl);
            }, true);

            buttonEl.addEventListener("click", (event) => {
                if (buttonEl.dataset.skipClickOnce === "1") {
                    event.preventDefault();
                    event.stopPropagation();
                    buttonEl.dataset.skipClickOnce = "0";
                }
            }, true);

            autoOpenOnce(buttonEl);
        }

        function tryMount() {
            attachDesktopBehavior();
        }

        tryMount();

        const observer = new MutationObserver(() => tryMount());
        observer.observe(rootDocument.body, {
            childList: true,
            subtree: true,
        });
        </script>
        """,
        height=0,
        width=0,
    )

# 这个字典专门给“任务卡”提供说明文字。
# key 统一用程序内部的任务 key，这样一级任务和二级任务都能共用。
TASK_CARD_DESCRIPTIONS = {
    "create": "从一个需求出发，生成传播提案、标题、导语或社媒文案。",
    "refine": "对已有稿件进行审稿、规范化或改写，让表达更贴近正式 PR 口径。",
    "generate": "适合从零开始生成传播提案、标题、导语或社媒文案。",
    "check": "适合检查品牌口径、时间表达、风险措辞和 PR 语气问题。",
    "normalize": "适合统一术语、品牌写法和更正式的官方表达。",
    "rewrite": "适合在不改变原意的前提下，把文本改得更像正式汽车 PR 稿件。",
}

def switch_task_selection(
    *,
    group_label: str | None = None,
    submode_label: str | None = None,
):
    request_navigation(
        action="switch_task",
        group_label=group_label,
        submode_label=submode_label,
    )

def clear_editor_state():
    st.session_state.last_result = ""
    st.session_state.main_input = ""
    st.session_state.main_input_area = ""
    st.session_state.has_unsaved_result = False


def execute_navigation(payload: dict):
    action = payload.get("action")

    if action == "go_home":
        st.session_state.has_selected_task = False
        st.session_state.task_group = None
        st.session_state.task_submode = None
        clear_editor_state()

    elif action == "switch_task":
        st.session_state.has_selected_task = True

        if payload.get("group_label") is not None:
            st.session_state.task_group = payload["group_label"]

        if payload.get("submode_label") is not None:
            st.session_state.task_submode = payload["submode_label"]

        clear_editor_state()

    st.session_state.pending_navigation = None
    st.rerun()


def request_navigation(action: str, group_label: str | None = None, submode_label: str | None = None):
    payload = {
        "action": action,
        "group_label": group_label,
        "submode_label": submode_label,
    }

    has_unsaved = bool(st.session_state.get("last_result", "").strip()) and st.session_state.get("has_unsaved_result", False)

    if has_unsaved:
        st.session_state.pending_navigation = payload
        st.rerun()
    else:
        execute_navigation(payload)

def go_back_to_task_home():
    request_navigation(action="go_home")

# 这个函数专门渲染“当前任务说明卡”。
# 目的：
# 1. 按钮保持短而清楚
# 2. 说明文字单独显示，不再挤进按钮
def render_task_summary_card(title: str, desc: str):
    st.markdown(
        f"""
        <div class="pr-task-summary-card">
            <div class="pr-task-summary-label">当前任务</div>
            <div class="pr-task-summary-title">{title}</div>
            <div class="pr-task-summary-desc">{desc}</div>
        </div>
        """,
        unsafe_allow_html=True
    )

def render_step_title(step_no: int, title: str, desc: str = ""):
    desc_html = f'<div class="pr-step-desc">{desc}</div>' if desc else ""

    st.markdown(
        f"""
        <div class="pr-step-title">{step_no}. {title}</div>
        {desc_html}
        """,
        unsafe_allow_html=True,
    )

def render_submitted_brief_card(title: str, body: str, source_name: str | None = None):
    source_html = ""
    if source_name:
        source_html = f"""
        <div style="font-size:12px; color:#6b7280; margin-top:8px;">
            已上传资料：{source_name}
        </div>
        """

    st.markdown(
        f"""
        <div style="
            border:1px solid rgba(15, 23, 42, 0.06);
            border-radius:18px;
            background:#fafafc;
            padding:16px 18px;
            margin-bottom:18px;
        ">
            <div style="font-size:12px; font-weight:700; color:#475467; margin-bottom:8px;">
                {title}
            </div>
            <div style="font-size:14px; line-height:1.75; color:#111827; white-space:pre-wrap;">
                {body}
            </div>
            {source_html}
        </div>
        """,
        unsafe_allow_html=True,
    )

def render_refine_source_summary_card(
    title: str,
    body: str,
    source_name: str | None = None,
):
    source_html = ""
    if source_name:
        source_html = f"""
        <div style="font-size:12px; color:#6b7280; margin-top:8px;">
            已上传稿件：{source_name}
        </div>
        """

    st.markdown(
        f"""
        <div style="
            border:1px solid rgba(15, 23, 42, 0.06);
            border-radius:18px;
            background:#fafafc;
            padding:16px 18px;
            margin-bottom:18px;
        ">
            <div style="font-size:12px; font-weight:700; color:#475467; margin-bottom:8px;">
                {title}
            </div>
            <div style="font-size:14px; line-height:1.75; color:#111827; white-space:pre-wrap;">
                {body}
            </div>
            {source_html}
        </div>
        """,
        unsafe_allow_html=True,
    )


def compact_preview(text: str, limit: int = 140) -> str:
    cleaned = re.sub(r"\s+", " ", (text or "").strip())
    if not cleaned:
        return "未填写"
    return cleaned if len(cleaned) <= limit else cleaned[:limit] + "…"


def render_processing_summary_card(
    *,
    mode: str,
    input_text: str = "",
    source_name: str | None = None,
    extra_term_rules: dict[str, str] | None = None,
    top_k: int | None = None,
):
    extra_term_rules = extra_term_rules or {}

    mode_name_map = {
        "generate": "生成新内容",
        "check": "审稿",
        "normalize": "规范化",
        "rewrite": "改写",
    }

    if mode == "generate" and not (input_text or "").strip() and source_name:
        input_text = "依据上传资料生成"

    if mode == "generate":
        input_label = "需求"
    else:
        input_label = "补充要求" if source_name else "待处理文本"

    source_text = source_name if source_name else "知识库"

    if extra_term_rules:
        terms_text = f"术语表 {len(extra_term_rules)} 条"
    else:
        terms_text = "无术语表"

    preview = compact_preview(input_text, limit=60)

    # 把所有信息拼成一行 tag 展示
    tags = [
        ("模式", mode_name_map.get(mode, mode)),
        ("来源", source_text),
    ]
    if mode == "generate" and top_k is not None:
        tags.append(("参考段落", f"Top {top_k}"))
    tags.append((input_label, preview))
    tags.append(("术语", terms_text))

    tags_html = ""
    for tag_label, tag_value in tags:
        tags_html += f"""
        <div style="
            display:inline-flex;
            align-items:center;
            gap:4px;
            background:#ffffff;
            border:1px solid rgba(15,23,42,0.07);
            border-radius:999px;
            padding:3px 10px;
            margin:2px 2px;
            font-size:12px;
            white-space:nowrap;
        ">
            <span style="color:#9ca3af;">{tag_label}</span>
            <span style="font-weight:600; color:#111827;">{html.escape(tag_value)}</span>
        </div>
        """

    st.markdown(
        f"""
        <div style="
            border:1px solid rgba(15, 23, 42, 0.06);
            border-radius:18px;
            background:#fafafc;
            padding:12px 16px;
            margin:4px 0 16px 0;
        ">
            <div style="font-size:11px; font-weight:700; color:#9ca3af; 
                        letter-spacing:0.04em; margin-bottom:8px;">本次处理</div>
            <div style="display:flex; flex-wrap:wrap; align-items:center;">{tags_html}</div>
        </div>
        """,
        unsafe_allow_html=True,
    )

def render_result_header(step_no: int, show_tools: bool = False):
    st.markdown(
        f'<div class="pr-step-title" style="margin-bottom:0;">{step_no}. 输出结果</div>',
        unsafe_allow_html=True,
    )

    if show_tools:
        st.markdown("<div style='height: 10px;'></div>", unsafe_allow_html=True)
        tool_left, tool_right = st.columns([0.48, 0.52], gap="small")

        with tool_left:
            export_format = st.selectbox(
                "",
                ["Markdown (.md)", "Word (.docx)", "TXT (.txt)"],
                index=0,
                key="export_format",
                label_visibility="collapsed",
            )

        with tool_right:
            export_data, export_mime, export_filename = get_export_payload(
                st.session_state.last_result,
                export_format
            )

            did_export = st.download_button(
                label="导出结果",
                data=export_data,
                file_name=export_filename,
                mime=export_mime,
                use_container_width=True,
                key="download_result_button",
            )

            if did_export:
                st.session_state.has_unsaved_result = False

def render_home_task_card(badge: str, title: str, desc: str):
    st.markdown(
        f"""
        <div class="pr-task-choice-card">
            <div class="pr-task-choice-badge">{badge}</div>
            <div class="pr-task-choice-title">{title}</div>
            <div class="pr-task-choice-desc">{desc}</div>
        </div>
        """,
        unsafe_allow_html=True,
    )
REFINE_MODE_BADGES = {
    "check": "审稿",
    "normalize": "规范化",
    "rewrite": "改写",
}

def render_unsaved_result_guard():
    pending = st.session_state.get("pending_navigation")
    if not pending:
        return

    # 改成温和的琥珀色提示，而不是红色报错感
    st.markdown(
        """
        <div style="
            border:1px solid rgba(180, 130, 20, 0.15);
            background:#fffbeb;
            border-radius:16px;
            padding:14px 18px;
            margin:10px 0 16px 0;
            display:flex;
            align-items:flex-start;
            gap:12px;
        ">
            <div style="font-size:18px; line-height:1; margin-top:1px;">💾</div>
            <div>
                <div style="font-size:14px; font-weight:700; color:#92400e; margin-bottom:4px;">
                    结果还没保存
                </div>
                <div style="font-size:13px; line-height:1.65; color:#b45309;">
                    离开后内容会清空，建议先导出再继续。
                </div>
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    # 三个按钮视觉分级：
    # 主动作（导出）占更多空间，"仍然离开"缩小弱化，"取消"最小
    guard_col1, guard_col2, guard_col3 = st.columns([0.50, 0.30, 0.20], gap="small")

    with guard_col1:
        did_download = st.download_button(
            label="📥 导出结果",
            data=st.session_state.get("last_result", "").encode("utf-8"),
            file_name="pr_agent_unsaved_result.md",
            mime="text/markdown",
            use_container_width=True,
            key="guard_download_unsaved_result",
            type="primary",
        )
        if did_download:
            st.session_state.has_unsaved_result = False

    with guard_col2:
        if st.button(
            "不保存，离开",
            use_container_width=True,
            key="confirm_leave_anyway",
        ):
            execute_navigation(pending)

    with guard_col3:
        if st.button(
            "取消",
            use_container_width=True,
            key="cancel_pending_navigation",
        ):
            st.session_state.pending_navigation = None
            st.rerun()

def build_refine_result_state_html(mode: str, title: str, desc: str) -> str:
    badge = REFINE_MODE_BADGES.get(mode, "处理")

    return f"""
    <div class="pr-result-state-card">
        <div class="pr-result-mode-badge">当前模式：{badge}</div>
        <div class="pr-result-state-title">{title}</div>
        <div class="pr-result-state-desc">{desc}</div>
    </div>
    """

 # 这个字典专门控制主按钮文案
 # 目的：不要再统一显示“运行当前任务”，而要让按钮直接表达动作。
RUN_BUTTON_LABELS = {
    "generate": "生成内容",
    "check": "开始审稿",
    "normalize": "开始规范化",
    "rewrite": "开始改写",
}

def get_example_text(mode: str) -> str:
    return MODE_INPUT_CONFIG.get(mode, {}).get("sample", "")

def call_with_supported_kwargs(func, *args, **kwargs):
    sig = inspect.signature(func)
    supported_kwargs = {
        key: value for key, value in kwargs.items()
        if key in sig.parameters
    }
    return func(*args, **supported_kwargs)

def sync_main_input():
    st.session_state.main_input = st.session_state.main_input_area

def normalize_column_name(name: str) -> str:
    return (
        str(name)
        .strip()
        .lower()
        .replace(" ", "")
        .replace("_", "")
        .replace("-", "")
    )


def parse_terms_table(uploaded_file) -> dict[str, str]:
    suffix = Path(uploaded_file.name).suffix.lower()

    if suffix == ".csv":
        df = pd.read_csv(uploaded_file)
    elif suffix in [".xlsx", ".xls"]:
        df = pd.read_excel(uploaded_file)
    else:
        raise ValueError("术语表仅支持 csv / xlsx / xls")

    if df.empty:
        raise ValueError("术语表为空")

    df = df.dropna(how="all").copy()
    df.columns = [str(col).strip() for col in df.columns]

    normalized_columns = {
        normalize_column_name(col): col for col in df.columns
    }

    source_aliases = ["错误写法", "原写法", "原词", "术语", "待替换", "before", "source", "from"]
    target_aliases = ["标准写法", "规范写法", "官方写法", "目标写法", "替换后", "after", "target", "to"]

    source_col = None
    target_col = None

    for alias in source_aliases:
        if normalize_column_name(alias) in normalized_columns:
            source_col = normalized_columns[normalize_column_name(alias)]
            break

    for alias in target_aliases:
        if normalize_column_name(alias) in normalized_columns:
            target_col = normalized_columns[normalize_column_name(alias)]
            break

    if source_col is None or target_col is None:
        if len(df.columns) >= 2:
            source_col = df.columns[0]
            target_col = df.columns[1]
        else:
            raise ValueError("术语表至少需要两列：原写法 / 标准写法")

    rules: dict[str, str] = {}
    for _, row in df[[source_col, target_col]].dropna().iterrows():
        raw_source = str(row[source_col]).strip()
        raw_target = str(row[target_col]).strip()

        if raw_source and raw_target:
            rules[raw_source] = raw_target

    if not rules:
        raise ValueError("术语表中没有可用规则")

    return rules


def apply_extra_term_rules(text: str, extra_terms: dict[str, str] | None = None) -> str:
    extra_terms = extra_terms or {}
    if not text or not extra_terms:
        return text

    updated_text = text
    for raw_source, raw_target in extra_terms.items():
        if not raw_source:
            continue
        updated_text = re.sub(re.escape(raw_source), raw_target, updated_text)

    return updated_text


def build_terms_prompt_block(extra_terms: dict[str, str] | None = None) -> str:
    extra_terms = extra_terms or {}
    if not extra_terms:
        return "【术语表】\n当前未上传术语表。"

    lines = ["【术语表】", "如涉及以下写法，请优先遵循："]
    for raw_source, raw_target in extra_terms.items():
        lines.append(f"- {raw_source} → {raw_target}")

    return "\n".join(lines)


@st.cache_resource
def get_markitdown():
    """
    缓存一个 MarkItDown 实例，避免每次 rerun 都重复初始化。
    """
    return MarkItDown(enable_plugins=False)


def clean_markdown_text(text: str) -> str:
    """
    对 MarkItDown 输出做一点轻量清洗，
    避免空行过多，保留基本结构。
    """
    if not text:
        return ""

    text = text.replace("\r\n", "\n").replace("\r", "\n")

    # 去掉连续过多空行
    while "\n\n\n" in text:
        text = text.replace("\n\n\n", "\n\n")

    return text.strip()


def read_uploaded_text(uploaded_file) -> str:
    """
    统一读取用户上传文件。
    - txt / md: 直接按文本读取
    - 其他常见文档格式: 用 MarkItDown 转成 Markdown，再作为纯文本送入后续流程
    """
    suffix = Path(uploaded_file.name).suffix.lower()
    file_bytes = uploaded_file.getvalue()

    # 1) 纯文本类，优先直接读，避免多走一层转换
    if suffix in [".txt", ".md"]:
        for encoding in ["utf-8", "utf-8-sig", "gb18030", "gbk"]:
            try:
                return file_bytes.decode(encoding).strip()
            except UnicodeDecodeError:
                continue
        raise ValueError("文本文件编码无法识别，建议保存为 UTF-8 后重试")

    # 2) 其余文档类，交给 MarkItDown
    supported_by_markitdown = {
        ".pdf", ".docx", ".pptx", ".xlsx", ".xls",
        ".csv", ".json", ".xml", ".html", ".htm"
    }

    if suffix in supported_by_markitdown:
        md = get_markitdown()

        # 保存到临时文件，再让 MarkItDown 按文件后缀识别
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
            tmp.write(file_bytes)
            tmp_path = tmp.name

        try:
            result = md.convert(tmp_path)
            text = getattr(result, "text_content", "") or ""
            text = clean_markdown_text(text)

            if not text.strip():
                raise ValueError("文件已读取，但未提取到可用文本内容")

            return text
        finally:
            try:
                Path(tmp_path).unlink(missing_ok=True)
            except Exception:
                pass

    raise ValueError(
        "暂不支持的文件类型。当前支持：txt, md, docx, pdf, pptx, xlsx, xls, csv, json, xml, html"
    )

# ==============================
# Helpers: generation with uploaded text
# ==============================

def build_uploaded_context(text: str, source_name: str) -> str:
    text = text.strip()
    if not text:
        return "未读取到上传资料内容。"
    return f"[上传资料 | 来源: {source_name}]\n{text}"


def generate_pr_from_uploaded_text(
    query: str,
    uploaded_text: str,
    source_name: str,
    extra_terms: dict[str, str] | None = None,
) -> str:
    extra_terms = extra_terms or {}

    query = apply_extra_term_rules(query, extra_terms)
    uploaded_text = apply_extra_term_rules(uploaded_text, extra_terms)

    context = build_uploaded_context(uploaded_text, source_name)
    terms_block = build_terms_prompt_block(extra_terms)

    core_prompt = f"""
你是一名服务汽车品牌的中国公关编辑，正在基于内部资料撰写传播提案内容。

请基于以下资料完成输出。

【重要要求】
1. 只能使用资料中明确出现的信息，不要补充资料中没有出现的配置、参数、功能和结论。
2. “传播方向”必须是传播切口，不要写成“空间/智能/豪华/操控”这类卖点分类。
3. 不要输出任何“媒体标题”内容。
4. 严格按照指定结构输出，不要增加新栏目。
5. 避免使用空泛AI套话。
6. 如涉及术语，请优先遵循术语表。

{terms_block}

【你只能输出以下四个部分，且必须按此顺序】
传播主题：
传播方向（3条）：
新闻稿段落：
社媒文案（3条）：

【参考资料】
{context}

【用户需求】
{query}
""".strip()

    title_prompt = f"""
你是一名中国汽车媒体编辑，请基于资料为一个汽车传播项目生成媒体标题。

【重要要求】
1. 只能使用资料中明确出现的信息。
2. 标题要更像传播选题，不要写成参数总结标题。
3. 如涉及术语，请优先遵循术语表。
4. 只允许输出编号标题，不要输出“传播主题”“媒体标题”“说明”等额外字段。

{terms_block}

【参考资料】
{context}

【用户需求】
{query}

请直接输出 5 个标题，每行一条，格式如下：
1.
2.
3.
4.
5.
""".strip()

    core_result = call_llm(core_prompt)
    core_result = strip_titles_section(core_result)

    title_result = call_llm(title_prompt)
    title_result = clean_title_result(title_result)

    final_result = f"{core_result}\n\n媒体标题（5条）：\n{title_result}"
    return final_result

def build_strategy_prompt(
    query: str,
    extra_terms: dict[str, str] | None = None,
    uploaded_text: str | None = None,
    source_name: str | None = None,
) -> str:
    extra_terms = extra_terms or {}
    query = apply_extra_term_rules(query, extra_terms)
    terms_block = build_terms_prompt_block(extra_terms)

    if uploaded_text and source_name:
        context_block = build_uploaded_context(
            apply_extra_term_rules(uploaded_text, extra_terms),
            source_name
        )
        source_rule = "可结合资料中的明确信息做传播判断，但不要补充资料中没有出现的事实。"
    else:
        context_block = "当前未上传参考资料。请仅基于用户需求完成传播判断，不要臆造具体配置、参数、功能和结论。"
        source_rule = "没有资料时，只做传播逻辑判断，不要写具体产品事实。"

    return f"""
你是一名服务汽车品牌的中国公关策略编辑。
现在不要直接写最终稿件，先完成传播策略判断。

【目标】
请先把这次任务的传播逻辑想清楚，再为后续生成具体交付内容做准备。

【重要要求】
1. 先做策略判断，不要直接输出完整新闻稿或媒体标题。
2. "传播方向"必须是传播切口，不要写成"空间/智能/豪华/操控"这类卖点分类。
3. 表达要像真实汽车 PR 提案前期判断，避免空泛AI套话。
4. {source_rule}
5. 如涉及术语，请优先遵循术语表。
6. 如果用户需求里有明确数量（如"三个传播主题"），传播方向建议的条数必须和用户要求一致。

{terms_block}

【参考信息】
{context_block}

【用户需求】
{query}

【请严格按以下结构输出】
任务理解：
传播对象判断：
核心传播命题：
传播方向建议（按用户需求数量输出）：
不建议采用的表达：
写作提醒：
""".strip()


def generate_pr_two_stage(
    query: str,
    top_k: int = 5,
    extra_terms: dict[str, str] | None = None,
    uploaded_text: str | None = None,
    source_name: str | None = None,
) -> str:
    extra_terms = extra_terms or {}
    query = apply_extra_term_rules(query, extra_terms)

    strategy_prompt = build_strategy_prompt(
        query=query,
        extra_terms=extra_terms,
        uploaded_text=uploaded_text,
        source_name=source_name,
    )
    strategy_result = call_llm(strategy_prompt).strip()

    if uploaded_text and source_name:
        delivery_result = generate_pr_from_uploaded_text(
            query=query,
            uploaded_text=uploaded_text,
            source_name=source_name,
            extra_terms=extra_terms,
        )
    else:
        try:
            core_result = call_with_supported_kwargs(
                generate_core_pr,
                query,
                top_k=top_k,
            )
            core_result = strip_titles_section(core_result)

            title_result = call_with_supported_kwargs(
                generate_titles,
                query,
                top_k=top_k,
            )
            title_result = clean_title_result(title_result)

            delivery_result = f"{core_result}\n\n媒体标题（5条）：\n{title_result}"

        except TypeError:
            fallback_result = call_with_supported_kwargs(
                generate_pr,
                query,
                top_k=top_k,
            )
            delivery_result = fallback_result

    return f"""【第一步｜传播策略骨架】
{strategy_result}



【第二步｜具体交付内容】
{delivery_result}"""

GENERATE_SECTION_HEADERS = {
    "angles": "传播方向（3条）：",
    "titles": "媒体标题（5条）：",
    "social": "社媒文案（3条）：",
}

ALL_GENERATE_HEADERS = [
    "传播主题：",
    "传播方向（3条）：",
    "新闻稿段落：",
    "社媒文案（3条）：",
    "媒体标题（5条）：",
]


def get_delivery_text(result_text: str) -> str:
    stage1, stage2 = split_two_stage_result(result_text)
    return (stage2 or result_text or "").strip()


def extract_section_content(text: str, header: str) -> str:
    joined = "|".join(re.escape(h) for h in ALL_GENERATE_HEADERS)
    pattern = rf"{re.escape(header)}\s*\n?(.*?)(?=\n(?:{joined})|\Z)"
    match = re.search(pattern, text, flags=re.DOTALL)
    return match.group(1).strip() if match else ""


def replace_section_content(text: str, header: str, new_content: str) -> str:
    joined = "|".join(re.escape(h) for h in ALL_GENERATE_HEADERS)
    pattern = rf"{re.escape(header)}\s*\n?(.*?)(?=\n(?:{joined})|\Z)"
    replacement = f"{header}\n{new_content.strip()}\n"
    if re.search(pattern, text, flags=re.DOTALL):
        return re.sub(pattern, replacement, text, count=1, flags=re.DOTALL).strip()
    return (text.rstrip() + "\n\n" + replacement).strip()


def regenerate_generate_section(
    *,
    section_key: str,
    query: str,
    current_result: str,
    top_k: int = 5,
    extra_terms: dict[str, str] | None = None,
    uploaded_text: str | None = None,
    source_name: str | None = None,
) -> str:
    header = GENERATE_SECTION_HEADERS[section_key]

    candidate_full = generate_pr_two_stage(
        query=query,
        top_k=top_k,
        extra_terms=extra_terms,
        uploaded_text=uploaded_text,
        source_name=source_name,
    )
    candidate_delivery = get_delivery_text(candidate_full)
    new_content = extract_section_content(candidate_delivery, header)

    if not new_content:
        raise ValueError("未能提取到对应内容，请重试。")

    current_stage1, current_delivery = split_two_stage_result(current_result)
    current_delivery = current_delivery or get_delivery_text(current_result)
    updated_delivery = replace_section_content(current_delivery, header, new_content)

    if current_stage1 is None:
        return updated_delivery

    return f"""【第一步｜传播策略骨架】
{current_stage1}



【第二步｜具体交付内容】
{updated_delivery}"""


def refine_existing_result(
    *,
    mode: str,
    current_result: str,
    instruction: str,
) -> str:
    if mode == "generate":
        prompt = f"""
你是一名服务汽车品牌的中国公关编辑。请基于用户补充要求，对已有结果做进一步优化。

【要求】
1. 保留原有整体结构，尤其保留：
- 【第一步｜传播策略骨架】
- 【第二步｜具体交付内容】
2. 第二步中的栏目标题尽量保持不变。
3. 不要新增没有依据的事实。
4. 优化方向只围绕用户补充要求展开。

【补充要求】
{instruction}

【当前结果】
{current_result}

请直接输出更新后的完整结果。
""".strip()
    else:
        prompt = f"""
你是一名服务汽车品牌的中国公关编辑。请基于用户补充要求，对已有处理结果做进一步优化。

【要求】
1. 保留原意，不新增事实。
2. 尽量保持现有结构。
3. 如果当前结果包含“问题清单 / 修改建议 / 参考改法”等结构，请尽量保留。
4. 只围绕用户补充要求优化表达。

【当前模式】
{mode}

【补充要求】
{instruction}

【当前结果】
{current_result}

请直接输出更新后的完整结果。
""".strip()

    return call_llm(prompt).strip()

# ==============================
# Helpers: save + rebuild index
# ==============================

def run_refine_with_uploaded_text(
    mode: str,
    source_text: str,
    instruction_text: str = "",
    extra_terms: dict[str, str] | None = None,
) -> str:
    extra_terms = extra_terms or {}
    source_text = apply_extra_term_rules(source_text, extra_terms)
    instruction_text = apply_extra_term_rules(instruction_text.strip(), extra_terms)

    mode_task_map = {
        "check": "请输出问题清单、修改建议与可参考改法。",
        "normalize": "请输出统一后的正式表达。",
        "rewrite": "请输出改写后的正式汽车 PR 风格文本。",
    }

    mode_rule_map = {
        "check": """
你是一名汽车品牌公关审稿编辑。
请基于上传稿件完成审稿，不要凭空补充稿件中没有的信息。
重点检查：
1. 品牌口径是否统一
2. 时间、事实、术语写法是否准确
3. 风险措辞是否需要收敛
4. 语气是否符合正式 PR 稿件
5. 结构与表达是否顺畅
""",
        "normalize": """
你是一名汽车品牌公关编辑。
请基于上传稿件做术语和表达规范化。
要求：
1. 优先统一品牌、产品、技术名词写法
2. 调整为更正式、更官方的表达
3. 尽量保留原意和原有结构
4. 不要新增稿件中没有的事实
""",
        "rewrite": """
你是一名汽车品牌公关编辑。
请基于上传稿件完成正式风格改写。
要求：
1. 保留原意，不新增事实
2. 语言更像正式汽车 PR 稿件
3. 语气更克制、更专业
4. 保持信息顺序尽量稳定
""",
    }

    extra_instruction_block = instruction_text if instruction_text else "无额外要求。"

    prompt = f"""
{mode_rule_map.get(mode, "")}

【额外个性化要求】
{extra_instruction_block}

【原稿内容】
{source_text}

【输出要求】
{mode_task_map.get(mode, "")}
""".strip()

    return call_llm(prompt)

def save_uploaded_file(uploaded_file, save_dir: Path) -> Path:
    save_dir.mkdir(parents=True, exist_ok=True)
    save_path = save_dir / uploaded_file.name
    with open(save_path, "wb") as f:
        f.write(uploaded_file.getbuffer())
    return save_path


def rebuild_index() -> tuple[bool, str]:
    try:
        result = subprocess.run(
            ["python", "-m", "src.indexer"],
            cwd=str(PROJECT_ROOT),
            capture_output=True,
            text=True,
            check=False,
        )

        if result.returncode == 0:
            return True, result.stdout.strip() or "索引重建完成"
        else:
            err = result.stderr.strip() or result.stdout.strip() or "索引重建失败"
            return False, err
    except Exception as e:
        return False, str(e)


# ==============================
# Helpers: examples
# ==============================

def build_docx_bytes(text: str, title: str = "PR Agent 导出结果") -> bytes:
    doc = Document()
    doc.add_heading(title, level=1)

    # 按空行分段，导出到 Word 里会更自然
    blocks = [block.strip() for block in text.split("\n\n") if block.strip()]
    for block in blocks:
        doc.add_paragraph(block)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def get_export_payload(text: str, export_format: str):
    if export_format == "Word (.docx)":
        return build_docx_bytes(text), "application/vnd.openxmlformats-officedocument.wordprocessingml.document", "pr_agent_output.docx"

    if export_format == "Markdown (.md)":
        return text.encode("utf-8"), "text/markdown", "pr_agent_output.md"

    return text.encode("utf-8"), "text/plain", "pr_agent_output.txt"

import html
# 这一行放在文件顶部 import 区
# html.escape() 用来安全地处理文本，避免特殊字符把 HTML 结构弄乱


def format_result_text(text: str) -> str:
    lines = text.splitlines()
    cleaned_lines = []

    for raw_line in lines:
        line = raw_line.strip()

        if not line:
            cleaned_lines.append("")
            continue

        if line.startswith("【") and line.endswith("】"):
            if cleaned_lines and cleaned_lines[-1] != "":
                cleaned_lines.append("")
            cleaned_lines.append(line)
            cleaned_lines.append("")
            continue

        if line.startswith(("• ", "- ", "* ")):
            cleaned_lines.append(f"- {line[2:].strip()}")
            continue

        cleaned_lines.append(line)

    formatted = "\n".join(cleaned_lines)

    while "\n\n\n" in formatted:
        formatted = formatted.replace("\n\n\n", "\n\n")

    return formatted
def build_result_html(text: str) -> str:
    lines = text.splitlines()
    parts = ['<div class="pr-result-html">']
    in_list = False

    for raw_line in lines:
        line = raw_line.strip()

        if not line:
            if in_list:
                parts.append("</ul>")
                in_list = False
            continue

        if line.startswith("【") and line.endswith("】"):
            if in_list:
                parts.append("</ul>")
                in_list = False
            parts.append(f"<h3>{html.escape(line)}</h3>")
            continue

        if line.startswith("- "):
            if not in_list:
                parts.append("<ul>")
                in_list = True
            parts.append(f"<li>{html.escape(line[2:].strip())}</li>")
            continue

        if in_list:
            parts.append("</ul>")
            in_list = False

        parts.append(f"<p>{html.escape(line)}</p>")

    if in_list:
        parts.append("</ul>")

    parts.append("</div>")
    return "".join(parts)
def split_two_stage_result(text: str) -> tuple[str | None, str | None]:
    text = text.strip()

    if "【第一步｜传播策略骨架】" not in text or "【第二步｜具体交付内容】" not in text:
        return None, text

    parts = text.split("【第二步｜具体交付内容】", 1)
    stage1_raw = parts[0].replace("【第一步｜传播策略骨架】", "").strip()
    stage2_raw = parts[1].strip()

    return stage1_raw, stage2_raw

# ==============================
# Session State
# ==============================


if "usage_assistant_answer" not in st.session_state:
    st.session_state.usage_assistant_answer = ""

# main_input: 记住主输入框里的内容。
# 这样 rerun 之后，用户刚才输入的文字不会丢。
if "main_input" not in st.session_state:
    st.session_state.main_input = ""

# last_result: 记住上一次运行产生的结果。
# 右侧结果区就是靠它来判断“显示结果”还是“显示空状态“
if "last_result" not in st.session_state:
    st.session_state.last_result = ""

# last_mode: 记住上一次使用的任务模式
# 这样当用户切换任务时，我们可以决定要不要清空旧结果。
if "last_mode" not in st.session_state:
    st.session_state.last_mode = "generate"

# 不默认选任务
if "task_group" not in st.session_state:
    st.session_state.task_group = None

if "task_submode" not in st.session_state:
    st.session_state.task_submode = None

# 是否已经进入任务
if "has_selected_task" not in st.session_state:
    st.session_state.has_selected_task = False

if "has_unsaved_result" not in st.session_state:
    st.session_state.has_unsaved_result = False

if "pending_navigation" not in st.session_state:
    st.session_state.pending_navigation = None

if "last_generate_query" not in st.session_state:
    st.session_state.last_generate_query = ""

if "last_generate_uploaded_text" not in st.session_state:
    st.session_state.last_generate_uploaded_text = None

if "last_generate_source_name" not in st.session_state:
    st.session_state.last_generate_source_name = None

if "last_refine_source_text" not in st.session_state:
    st.session_state.last_refine_source_text = ""

if "last_refine_source_name" not in st.session_state:
    st.session_state.last_refine_source_name = None

# 记录上一次生成时参考了哪些知识库资料
if "last_search_sources" not in st.session_state:
    st.session_state.last_search_sources = []

# ==============================
# Derived Task State
# ==============================

current_group_key = None
current_available_modes = []
current_available_modes_labels = []
mode = None

# 只有真正选过任务后，才开始推导当前 mode
if st.session_state.has_selected_task and st.session_state.task_group is not None:
    current_group_key = TASK_GROUP_LABEL_TO_KEY[st.session_state.task_group]
    current_available_modes = GROUP_TO_MODES[current_group_key]
    current_available_modes_labels = [MODE_LABELS[m] for m in current_available_modes]

    # 如果当前保存的 task_submode 不在这一组任务里，就自动修正
    if st.session_state.task_submode not in current_available_modes_labels:
        st.session_state.task_submode = current_available_modes_labels[0]

    mode = LABEL_TO_MODE[st.session_state.task_submode]


# ==============================
# Header
# ==============================

PAGE_COLS = [0.06, 0.88, 0.06]

page_left, page_center, page_right = st.columns(PAGE_COLS)

with page_center:
    st.markdown("""
    <div class="pr-hero">
        <div class="pr-eyebrow">Automotive PR Workflow Assistant</div>
        <div class="pr-title">PR Agent</div>
        <div class="pr-subtitle">面向汽车公关场景的知识检索、内容生成、审稿与术语规范化工作台</div>
    </div>
    """, unsafe_allow_html=True)
    render_unsaved_result_guard()

# ==============================
# Sidebar
# ==============================

with st.sidebar:
    # 侧边栏现在只放“高级功能”
    # 也就是：第一次使用可以先不碰，但进阶用户会需要的内容。
    st.markdown("### 高级功能")
    
    # 这一句说明侧边栏是干什么的。
    # 我把它压缩短了，避免用户一打开就看到很多解释文字。
    st.caption("术语表、参考资料设置和知识库维护。第一次使用可以先忽略。")
    
    # 先给这些变量一个默认值。
    # 这样即使用户什么都没上传，后面的运行逻辑也不会报“变量未定义”。
    uploaded_terms = None
    extra_term_rules = {}
    top_k = 5
    kb_upload = None
    save_button = False
    rebuild_button = False

    # 上传术语表。
    # 作用：让当前会话里生成 / 审稿 / 规范化 /改写都能参考统一口径。
    uploaded_terms = st.file_uploader(
        "上传术语表（可选）",
        type=["xlsx", "csv"],
        accept_multiple_files=False,
        key="terms_uploader",
        help="当前会话有效，用于统一术语和官方表达",
        width="stretch",
    )

    # 如果用户上传了术语表，就尝试解析。
    if uploaded_terms is not None:
        try:
            # 解析结果会变成一个 dict，例如：
            # {"奔驰"：“梅赛德斯-奔驰”，“德国时间”：“欧洲中部时间[CET]"}
            extra_term_rules = parse_terms_table(uploaded_terms)

            # 成功后提示一共加载了多少条规则。
            st.success(f"术语表已加载：{len(extra_term_rules)}条规则")
        except Exception as e:
            # 如果术语表格式不对，就在侧边栏给出错误提示。
            st.error(f"术语表读取失败：{e}")

    # Top K 只在 generate 模式下有意义。
    # 因为它控制“从知识库里取多少段最相关内容”。
    top_k = 5
    if mode == "generate":
        top_k = st.number_input(
            "参考资料数量",
            min_value=1,
            max_value=10,
            value=5,
            help="生成内容时，系统会优先参考前 N 段最相关资料。数值越大，参考范围越广。",
        )
    # 加一条分割线，让“术语表”和“知识库维护”视觉上分开。
    st.markdown("---")
    st.caption("知识库维护")

    # 这里是把文件长期加入知识库，而不是只用于本次运行。
    kb_upload = st.file_uploader(
        "上传文件加入知识库",
        type=["txt", "md", "docx", "pdf", "pptx", "xlsx", "xls", "csv", "json", "xml", "html", "htm"],
        accept_multiple_files=False,
        key="kb_upload",
    )

    # 保存按钮：把上传文件保存到知识库目录。
    save_button = st.button("保存到知识库目录")

    # 重建索引按钮：保存后重新做向量索引，供后续检索使用。
    rebuild_button = st.button("重建索引")

# ==============================
# Knowledge Base Actions
# ==============================

if save_button:
    if kb_upload is None:
        st.warning("请先选择一个要加入知识库的文件")
    else:
        try:
            saved_path = save_uploaded_file(kb_upload, UPLOAD_DIR)
            st.success(f"文件已保存到：{saved_path}")
        except Exception as e:
            st.error(f"保存失败：{e}")

if rebuild_button:
    with st.spinner("正在重建索引..."):
        ok, msg = rebuild_index()
        if ok:
            st.success("索引重建完成")
            if msg:
                st.text_area("重建日志", value=msg, height=240)
        else:
            st.error(f"索引重建失败：{msg}")


show_result_panel = bool(st.session_state.last_result.strip())
# ==============================
# Step 1: Task Selection
# ==============================

# ------------------------------
# 状态 A：还没有进入任务
# 只显示一级任务入口
# ------------------------------

# ==============================
# Global Usage Assistant
# 桌面 App 风格：右下角悬浮机器人 + popover
# ==============================

assistant_mode = mode or "generate"
assistant_has_uploaded_file = False

if st.session_state.has_selected_task:
    if assistant_mode == "generate":
        assistant_has_uploaded_file = st.session_state.get("uploaded_material") is not None
    else:
        assistant_has_uploaded_file = st.session_state.get("uploaded_manuscript") is not None

with st.popover("🤖"):
    render_usage_assistant_content(
        mode=assistant_mode,
        has_uploaded_file=assistant_has_uploaded_file,
    )

inject_desktop_assistant_js()

if not st.session_state.has_selected_task:
    outer_left, center, outer_right = st.columns(PAGE_COLS)

    with center:
        st.markdown('<div class="pr-home-inner">', unsafe_allow_html=True)

        st.markdown(
            """
            <div class="pr-step-title" style="margin-bottom:4px;">选择任务</div>
            <div class="pr-step-desc" style="margin-bottom:18px;">
                先选择你当前要完成的 PR 工作，再进入输入与结果页面。
            </div>
            """,
            unsafe_allow_html=True,
        )

        group_col1, group_col2 = st.columns(2, gap="large")

        with group_col1:
            render_home_task_card(
                "CREATE",
                "生成新内容",
                "从一个需求出发，生成传播提案、标题、导语或社媒文案。",
            )
            st.markdown('<div class="pr-home-actions">', unsafe_allow_html=True)
            if st.button(
                "进入生成任务",
                key="task_group_create_button",
                use_container_width=True,
                type="secondary",
            ):
                switch_task_selection(
                    group_label=TASK_GROUP_LABELS["create"],
                    submode_label=MODE_LABELS["generate"],
                )
            st.markdown('</div>', unsafe_allow_html=True)

        with group_col2:
            render_home_task_card(
                "REFINE",
                "处理现有稿件",
                "对已有稿件进行审稿、规范化或改写。",
            )
            st.markdown('<div class="pr-home-actions">', unsafe_allow_html=True)
            if st.button(
                "进入处理任务",
                key="task_group_refine_button",
                use_container_width=True,
                type="secondary",
            ):
                switch_task_selection(
                    group_label=TASK_GROUP_LABELS["refine"],
                    submode_label=MODE_LABELS["check"],
                )
            st.markdown('</div>', unsafe_allow_html=True)

        st.markdown("</div>", unsafe_allow_html=True)

    st.stop()

# ------------------------------
# 状态 B：已经进入任务
# generate：单列主流程
# refine：左右对照
# ------------------------------

page_left, page_center, page_right = st.columns(PAGE_COLS)

SECTION_2_TITLES = {
    "generate": "输入传播需求",
    "check": "输入待审稿文本",
    "normalize": "输入待规范化文本",
    "rewrite": "输入待改写文本",
}

show_result_panel = bool(st.session_state.last_result.strip())
is_create_flow = st.session_state.task_group == TASK_GROUP_LABELS["create"]
is_generate_mode = mode == "generate"

input_step_no = 2 if is_create_flow else 3
result_step_no = 3 if is_create_flow else 4

uploaded_material = None
uploaded_manuscript = None
run_button = False
user_input = st.session_state.get("main_input_area", st.session_state.main_input)
partial_generate_section = None
followup_submit = False
followup_instruction = ""

with page_center:
    # 顶部返回
    back_col, _ = st.columns([0.26, 0.74], gap="small")
    with back_col:
        if st.button("← 返回任务选择", key="back_to_task_home", use_container_width=False):
            go_back_to_task_home()

    # refine 才显示步骤 2
    if not is_create_flow:
        st.markdown("<div style='height: 12px;'></div>", unsafe_allow_html=True)
        render_step_title(2, "选择处理方式", "")
        st.markdown("<div style='height: 10px;'></div>", unsafe_allow_html=True)

        sub_col1, sub_col2, sub_col3 = st.columns(3, gap="small")

        with sub_col1:
            check_selected = st.session_state.task_submode == MODE_LABELS["check"]
            if st.button(
                "审稿",
                key="task_submode_check_button",
                use_container_width=True,
                type="primary" if check_selected else "secondary",
            ):
                switch_task_selection(submode_label=MODE_LABELS["check"])

        with sub_col2:
            normalize_selected = st.session_state.task_submode == MODE_LABELS["normalize"]
            if st.button(
                "规范化",
                key="task_submode_normalize_button",
                use_container_width=True,
                type="primary" if normalize_selected else "secondary",
            ):
                switch_task_selection(submode_label=MODE_LABELS["normalize"])

        with sub_col3:
            rewrite_selected = st.session_state.task_submode == MODE_LABELS["rewrite"]
            if st.button(
                "改写",
                key="task_submode_rewrite_button",
                use_container_width=True,
                type="primary" if rewrite_selected else "secondary",
            ):
                switch_task_selection(submode_label=MODE_LABELS["rewrite"])

    st.markdown("<div style='height: 20px;'></div>", unsafe_allow_html=True)

    # =========================================================
    # A. generate：单列
    # =========================================================
    if is_generate_mode:
        # 生成前：展示完整输入区
        if not show_result_panel:
            render_step_title(input_step_no, SECTION_2_TITLES[mode], "")

            input_config = MODE_INPUT_CONFIG[mode]

            uploaded_material = st.file_uploader(
                "上传参考资料（可选）",
                type=["txt", "md", "docx", "pdf", "pptx", "xlsx", "xls", "csv", "json", "xml", "html", "htm"],
                help="支持 TXT / DOCX。可不上传，直接根据输入需求生成。",
                key="uploaded_material",
            )

            header_left, header_right = st.columns([0.84, 0.16], gap="small")

            with header_left:
                st.markdown(
                    f"""
                    <div style="font-size:14px; font-weight:600; color:#111827; margin-bottom:6px;">
                        {input_config["field_label"]}
                    </div>
                    """,
                    unsafe_allow_html=True,
                )

            with header_right:
                if st.button("示例", use_container_width=True, key="try_example_button"):
                    sample_text = input_config["sample"]
                    st.session_state.main_input = sample_text
                    st.session_state.main_input_area = sample_text
                    st.rerun()

            user_input = st.text_area(
                label="",
                placeholder=input_config["placeholder"],
                help=input_config["help"],
                height=260,
                key="main_input_area",
                on_change=sync_main_input,
                label_visibility="collapsed",
            )

            current_input_text = st.session_state.get("main_input_area", "")
            st.session_state.main_input = current_input_text
            user_input = current_input_text
            char_count = len(re.sub(r"\s+", "", current_input_text))

            meta_left, meta_right = st.columns([0.52, 0.48], gap="small")

            with meta_left:
                st.markdown(
                    f"""
                    <div style="font-size:12px; color:#6b7280; margin-top:4px; margin-bottom:12px;">
                        已提交 {char_count} 字
                        <span
                            title="输入后按 Cmd+Enter 更新内容与字数统计"
                            style="
                                display:inline-flex;
                                align-items:center;
                                justify-content:center;
                                width:16px;
                                height:16px;
                                margin-left:6px;
                                border-radius:999px;
                                border:1px solid rgba(15, 23, 42, 0.12);
                                color:#9ca3af;
                                font-size:11px;
                                cursor:help;
                                vertical-align:middle;
                            "
                        >?</span>
                    </div>
                    """,
                    unsafe_allow_html=True,
                )

            with meta_right:
                st.markdown(
                    f"""
                    <div style="font-size:12px; color:#6b7280; text-align:right; margin-top:4px; margin-bottom:12px;">
                        输出：{MODE_OUTPUT_HINTS[mode]}
                    </div>
                    """,
                    unsafe_allow_html=True,
                )

            uploaded_source_name = uploaded_material.name if uploaded_material is not None else None

            if current_input_text.strip() or uploaded_material is not None or extra_term_rules:
                render_processing_summary_card(
                    mode=mode,
                    input_text=current_input_text,
                    source_name=uploaded_source_name,
                    extra_term_rules=extra_term_rules,
                    top_k=top_k,
                )

            button_left, button_right = st.columns([0.28, 0.72], gap="medium")

            with button_left:
                if st.button("清空", use_container_width=True, key="clear_input_button"):
                    st.session_state.main_input = ""
                    st.session_state.main_input_area = ""
                    st.rerun()

            with button_right:
                run_button = st.button(
                    RUN_BUTTON_LABELS[mode],
                    type="primary",
                    use_container_width=True,
                    key="run_main_button",
                )

            run_status_placeholder = st.empty()
            run_progress_placeholder = st.empty()

        # 生成后：输入缩成已提交需求卡，结果成为主内容
        else:
            render_result_header(result_step_no, show_tools=True)

            st.markdown("<div style='height: 14px;'></div>", unsafe_allow_html=True)

            current_input_text = st.session_state.get("main_input_area", "").strip()
            uploaded_source_name = None
            if "uploaded_material" in st.session_state and st.session_state.uploaded_material is not None:
                try:
                    uploaded_source_name = st.session_state.uploaded_material.name
                except Exception:
                    uploaded_source_name = None

            render_submitted_brief_card(
                "已提交需求",
                current_input_text or "（本次未保留输入文本）",
                uploaded_source_name,
            )

            stage1_raw, stage2_raw = split_two_stage_result(st.session_state.last_result)

            if stage1_raw is None:
                with st.expander("查看结果", expanded=True):
                    st.markdown(
                        build_result_html(format_result_text(stage2_raw or "")),
                        unsafe_allow_html=True,
                    )
            else:
                with st.expander("STEP 1｜传播策略骨架", expanded=True):
                    st.markdown(
                        build_result_html(format_result_text(stage1_raw)),
                        unsafe_allow_html=True,
                    )

                with st.expander("STEP 2｜具体交付内容", expanded=True):
                    st.markdown(
                        build_result_html(format_result_text(stage2_raw)),
                        unsafe_allow_html=True,
                    )

            st.markdown("<div style='height: 18px;'></div>", unsafe_allow_html=True)

            run_status_placeholder = st.empty()
            run_progress_placeholder = st.empty()

            action_left, action_right = st.columns([0.26, 0.74], gap="medium")

            with action_left:
                if st.button("重新编辑需求", use_container_width=True, key="edit_generate_input_button"):
                    st.session_state.last_result = ""
                    st.session_state.has_unsaved_result = False
                    st.rerun()

            with action_right:
                run_button = st.button(
                    "基于当前需求重新生成",
                    type="primary",
                    use_container_width=True,
                    key="rerun_generate_button",
                )

            st.markdown("<div style='height: 10px;'></div>", unsafe_allow_html=True)

            partial_col1, partial_col2, partial_col3 = st.columns(3, gap="small")

            with partial_col1:
                if st.button("只重生成传播方向", use_container_width=True, key="rerun_angles_button"):
                    partial_generate_section = "angles"

            with partial_col2:
                if st.button("只重生成标题", use_container_width=True, key="rerun_titles_button"):
                    partial_generate_section = "titles"

            with partial_col3:
                if st.button("只重生成社媒文案", use_container_width=True, key="rerun_social_button"):
                    partial_generate_section = "social"

            
    # =========================================================
    # B. refine：无结果前左右对照；有结果后左摘要右结果
    # =========================================================
    else:
        # -----------------------------
        # refine：还没出结果
        # -----------------------------
        if not show_result_panel:
            left, right = st.columns([1.08, 0.92], gap="large")

            with left:
                render_step_title(input_step_no, SECTION_2_TITLES[mode], "")

                input_config = MODE_INPUT_CONFIG[mode]

                uploaded_manuscript = st.file_uploader(
                    "上传稿件（可选）",
                    type=["txt", "md", "docx", "pdf", "pptx", "xlsx", "xls", "csv", "json", "xml", "html", "htm"],
                    help="支持 TXT / DOCX。也可直接粘贴文本进行处理。",
                    key="uploaded_manuscript",
                )

                input_config = get_refine_input_config(
                    mode,
                    uploaded_manuscript is not None,
                )

                if uploaded_manuscript is not None:
                    st.markdown(
                        """
                        <div style="
                            font-size:13px;
                            color:#6b7280;
                            line-height:1.7;
                            margin: 2px 0 10px 0;
                        ">
                            已上传稿件。下面可以补充个性化处理要求，例如：重点检查品牌口径；不要整段重写；只统一术语。
                        </div>
                        """,
                        unsafe_allow_html=True,
                    )

                header_left, header_right = st.columns([0.78, 0.22], gap="small")

                with header_left:
                    st.markdown(
                        f"""
                        <div style="font-size:14px; font-weight:600; color:#111827; margin-bottom:6px;">
                            {input_config["field_label"]}
                        </div>
                        """,
                        unsafe_allow_html=True,
                    )

                with header_right:
                    if st.button("示例", use_container_width=False, key="try_example_button"):
                        sample_text = input_config["sample"]
                        st.session_state.main_input = sample_text
                        st.session_state.main_input_area = sample_text
                        st.rerun()

                user_input = st.text_area(
                    label="",
                    placeholder=input_config["placeholder"],
                    help=input_config["help"],
                    height=260,
                    key="main_input_area",
                    on_change=sync_main_input,
                    label_visibility="collapsed",
                )

                current_input_text = st.session_state.get("main_input_area", "")
                st.session_state.main_input = current_input_text
                user_input = current_input_text
                char_count = len(re.sub(r"\s+", "", current_input_text))

                meta_left, meta_right = st.columns([0.52, 0.48], gap="small")

                with meta_left:
                    st.markdown(
                        f"""
                        <div style="font-size:12px; color:#6b7280; margin-top:4px; margin-bottom:12px;">
                            已提交 {char_count} 字
                        </div>
                        """,
                        unsafe_allow_html=True,
                    )

                with meta_right:
                    st.markdown(
                        f"""
                        <div style="font-size:12px; color:#6b7280; text-align:right; margin-top:4px; margin-bottom:12px;">
                            输出：{MODE_OUTPUT_HINTS[mode]}
                        </div>
                        """,
                        unsafe_allow_html=True,
                    )
                uploaded_source_name = uploaded_manuscript.name if uploaded_manuscript is not None else None

                if current_input_text.strip() or uploaded_manuscript is not None or extra_term_rules:
                    render_processing_summary_card(
                        mode=mode,
                        input_text=current_input_text,
                        source_name=uploaded_source_name,
                        extra_term_rules=extra_term_rules,
                        top_k=None,
                    )

                button_left, button_right = st.columns([0.28, 0.72], gap="medium")

                with button_left:
                    if st.button("清空", use_container_width=True, key="clear_input_button"):
                        st.session_state.main_input = ""
                        st.session_state.main_input_area = ""
                        st.rerun()

                with button_right:
                    run_button = st.button(
                        RUN_BUTTON_LABELS[mode],
                        type="primary",
                        use_container_width=True,
                        key="run_main_button",
                    )

            with right:
                render_result_header(result_step_no, show_tools=False)
                st.markdown("<div style='height: 12px;'></div>", unsafe_allow_html=True)

                result_state_placeholder = st.empty()
                run_status_placeholder = st.empty()
                run_progress_placeholder = st.empty()

                empty_desc_map = {
                    "check": "运行后在这里查看问题清单、修改建议与参考改法。",
                    "normalize": "运行后在这里查看统一后的术语与正式表达。",
                    "rewrite": "运行后在这里查看改写后的正式 PR 风格文本。",
                }

                result_state_placeholder.markdown(
                    build_refine_result_state_html(
                        mode,
                        "结果面板",
                        empty_desc_map.get(mode, "运行后在这里查看处理结果。"),
                    ),
                    unsafe_allow_html=True,
                )

    # -----------------------------
    # refine：已经出结果
    # -----------------------------
        else:
            render_result_header(result_step_no, show_tools=True)
            st.markdown("<div style='height: 14px;'></div>", unsafe_allow_html=True)

            stage1_raw, stage2_raw = split_two_stage_result(st.session_state.last_result)
            processed_text = stage2_raw or st.session_state.last_result

            summary_col, result_col = st.columns([0.34, 0.66], gap="large")

            with summary_col:
                current_input_text = st.session_state.get("main_input_area", "").strip()

                uploaded_source_name = None
                if "uploaded_manuscript" in st.session_state and st.session_state.uploaded_manuscript is not None:
                    try:
                        uploaded_source_name = st.session_state.uploaded_manuscript.name
                    except Exception:
                        uploaded_source_name = None

                render_refine_source_summary_card(
                    "已提交来源",
                    current_input_text or "（本次主要基于上传稿件处理）",
                    uploaded_source_name,
                )

                if st.session_state.get("last_refine_source_text"):
                    with st.expander("原文 / 处理后对照", expanded=False):
                        compare_left, compare_right = st.columns(2, gap="small")

                        with compare_left:
                            st.text_area(
                                "原文",
                                value=st.session_state.get("last_refine_source_text", ""),
                                height=220,
                                disabled=True,
                                key=f"refine_original_compare_{mode}",
                            )

                        with compare_right:
                            st.text_area(
                                "处理后",
                                value=processed_text,
                                height=220,
                                disabled=True,
                                key=f"refine_result_compare_{mode}",
                            )

                st.markdown("<div style='height: 8px;'></div>", unsafe_allow_html=True)

                followup_instruction = st.text_input(
                    "继续细化当前结果",
                    placeholder="例如：语气更克制；不要改动段落结构；保留问题清单，只把建议写得更简洁",
                    key=f"refine_followup_instruction_{mode}",
                )

                followup_submit = st.button(
                    "应用到当前结果",
                    use_container_width=True,
                    key=f"apply_refine_followup_button_{mode}",
                )

                st.markdown("<div style='height: 8px;'></div>", unsafe_allow_html=True)

                if st.button("重新编辑文本", use_container_width=True, key="edit_refine_input_button"):
                    st.session_state.last_result = ""
                    st.rerun()

            with result_col:
                done_desc_map = {
                    "check": "已生成问题清单、修改建议与参考改法。",
                    "normalize": "已统一术语、品牌写法与更正式的官方表达。",
                    "rewrite": "已输出更贴近正式汽车 PR 风格的改写结果。",
                }

                st.markdown(
                    build_refine_result_state_html(
                        mode,
                        "处理完成",
                        done_desc_map.get(mode, "结果已生成。"),
                    ),
                    unsafe_allow_html=True,
                )

                run_status_placeholder = st.empty()
                run_progress_placeholder = st.empty()

                st.markdown('<div class="pr-result-scroll">', unsafe_allow_html=True)
                # 【新加】展示知识库参考来源
                sources = st.session_state.get("last_search_sources", [])
                if sources and mode == "generate":
                    with st.expander(f"📚 知识库参考来源（共 {len(sources)} 条）", expanded=False):
                        for i, src in enumerate(sources, 1):
                            file_name = src.get("file_name", "未知文件")
                            car_model = src.get("car_model") or "通用"
                            doc_type = src.get("doc_type", "")
                            score = src.get("rerank_score", 0)
                            breakdown = src.get("score_breakdown", {})
                            text_preview = src.get("text", "")[:120]

                            doc_type_label = {
                                "official_pr": "正式新闻稿",
                                "product_material": "产品资料",
                                "subtitle": "字幕稿",
                                "meeting_note": "会议纪要",
                                "other": "其他资料",
                            }.get(doc_type, doc_type)

                            st.markdown(
                                f"""
                                <div style="
                                    border:1px solid rgba(15,23,42,0.06);
                                    border-radius:14px;
                                    background:#fafafc;
                                    padding:12px 14px;
                                    margin-bottom:10px;
                                ">
                                    <div style="font-size:13px; font-weight:700; color:#111827; margin-bottom:4px;">
                                        {i}. {html.escape(file_name)}
                                    </div>
                                    <div style="font-size:12px; color:#6b7280; margin-bottom:6px;">
                                        车型：{car_model} ｜ 类型：{doc_type_label} ｜ 综合评分：{score:.3f}
                                    </div>
                                    <div style="font-size:13px; color:#374151; line-height:1.65;">
                                        {html.escape(text_preview)}{"…" if len(src.get("text","")) > 120 else ""}
                                    </div>
                                </div>
                                """,
                                unsafe_allow_html=True,
                            )
                if stage1_raw is None:
                    with st.expander("查看结果", expanded=True):
                        st.markdown(
                            build_result_html(format_result_text(processed_text)),
                            unsafe_allow_html=True,
                        )
                else:
                    with st.expander("STEP 1｜传播策略骨架", expanded=True):
                        st.markdown(
                            build_result_html(format_result_text(stage1_raw)),
                            unsafe_allow_html=True,
                        )

                    with st.expander("STEP 2｜具体交付内容", expanded=True):
                        st.markdown(
                            build_result_html(format_result_text(stage2_raw)),
                            unsafe_allow_html=True,
                        )

                st.markdown("</div>", unsafe_allow_html=True)
# ==============================
# Run
# ==============================

if run_button or partial_generate_section is not None or followup_submit:
    if followup_submit and not followup_instruction.strip():
        st.warning("请先输入继续细化要求")
        st.stop()

    if partial_generate_section is not None:
        base_query = st.session_state.get("last_generate_query", "").strip()
        base_uploaded_text = st.session_state.get("last_generate_uploaded_text")
        if not base_query and not base_uploaded_text:
            st.warning("缺少可用于局部重生成的原始需求，请先完整生成一次")
            st.stop()

    if partial_generate_section is None and not followup_submit:
        if mode == "generate":
            if not user_input.strip() and uploaded_material is None:
                st.warning("请先输入传播需求，或上传参考资料")
                st.stop()
        else:
            if uploaded_manuscript is None and not user_input.strip():
                st.warning("请先输入待处理文本，或上传稿件")
                st.stop()

    try:
        result = ""
        progress_bar = run_progress_placeholder.progress(0)
        from src.search import search as kb_search   
        if followup_submit:
            run_status_placeholder.markdown(
                '<div class="pr-run-status">正在根据补充要求优化当前结果...</div>',
                unsafe_allow_html=True,
            )
        elif partial_generate_section is not None:
            section_name_map = {
                "angles": "传播方向",
                "titles": "标题",
                "social": "社媒文案",
            }
            run_status_placeholder.markdown(
                f'<div class="pr-run-status">正在重生成{section_name_map.get(partial_generate_section, "")}...</div>',
                unsafe_allow_html=True,
            )

        if mode == "generate":
            progress_bar.progress(25)

            if uploaded_material is not None:
                material_text = read_uploaded_text(uploaded_material)
                if not material_text.strip():
                    raise ValueError("上传资料为空或无法读取内容")

                st.session_state.last_generate_query = user_input.strip()
                st.session_state.last_generate_uploaded_text = material_text
                st.session_state.last_generate_source_name = uploaded_material.name

                progress_bar.progress(45)

                # 【新加】搜索知识库，把来源存起来，界面上展示给用户
                _sources = kb_search(user_input.strip() or uploaded_material.name, k=top_k)
                st.session_state.last_search_sources = _sources

                result = generate_pr_two_stage(
                    query=user_input,
                    top_k=top_k,
                    extra_terms=extra_term_rules,
                    uploaded_text=material_text,
                    source_name=uploaded_material.name,
                )

                progress_bar.progress(85)

            else:
                st.session_state.last_generate_query = user_input.strip()
                st.session_state.last_generate_uploaded_text = None
                st.session_state.last_generate_source_name = None

                progress_bar.progress(45)

                # 【新加】搜索知识库，把来源存起来，界面上展示给用户
                _sources = kb_search(user_input.strip(), k=top_k)
                st.session_state.last_search_sources = _sources

                result = generate_pr_two_stage(
                    query=user_input,
                    top_k=top_k,
                    extra_terms=extra_term_rules,
                )

                progress_bar.progress(85)

        else:
            instruction_text = user_input.strip()
            progress_bar.progress(25)

            if uploaded_manuscript is not None:
                manuscript_text = read_uploaded_text(uploaded_manuscript)
                if not manuscript_text.strip():
                    raise ValueError("上传稿件为空或无法读取内容")

                st.session_state.last_refine_source_text = manuscript_text
                st.session_state.last_refine_source_name = uploaded_manuscript.name

                progress_bar.progress(45)

                result = run_refine_with_uploaded_text(
                    mode=mode,
                    source_text=manuscript_text,
                    instruction_text=instruction_text,
                    extra_terms=extra_term_rules,
                )

                progress_bar.progress(85)

            else:
                input_text = user_input.strip()

                progress_bar.progress(45)

                if not input_text:
                    raise ValueError("没有可处理的文本内容")

                st.session_state.last_refine_source_text = input_text
                st.session_state.last_refine_source_name = None

                input_text = apply_extra_term_rules(input_text, extra_term_rules)
                progress_bar.progress(65)

                if mode == "check":
                    result = check_text(input_text)
                elif mode == "normalize":
                    result = normalize_pr(input_text)
                elif mode == "rewrite":
                    result = rewrite_text(input_text)
                else:
                    raise ValueError("未知模式")

                progress_bar.progress(85)

        progress_bar.progress(100)
        st.session_state.last_result = result
        st.session_state.has_unsaved_result = True

        if followup_submit:
            if mode == "generate":
                st.session_state.generate_followup_instruction = ""
            else:
                st.session_state[f"refine_followup_instruction_{mode}"] = ""

        st.rerun()

    except Exception as e:
        run_progress_placeholder.empty()
        run_status_placeholder.empty()
        st.error(f"运行失败：{e}")