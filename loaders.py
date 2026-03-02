from pathlib import Path
from docx import Document


def read_docx(path: Path) -> str:
    doc = Document(str(path))
    texts = []

    # 方式1：普通段落
    for p in doc.paragraphs:
        if p.text.strip():
            texts.append(p.text.strip())

    # 方式2：表格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    texts.append(cell.text.strip())

    # 方式3：遍历所有 XML 节点（强力模式）
    for element in doc.element.body.iter():
        if element.text and element.text.strip():
            texts.append(element.text.strip())

    # 去重
    texts = list(dict.fromkeys(texts))

    return "\n".join(texts)